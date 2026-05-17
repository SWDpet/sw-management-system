"""
점검내역서 시안 3종 (.docx) 생성기 — 시안 A 클래식 / B 모던 / C 미니멀.

요건:
  • 표지: 사업명 / 월 / 기관(지자체·고객) / 확인자(담당자=공무원·고객) 소속·성명·서명란 / 점검자 동일
  • 2페이지: 월별 점검사항 표 (커스터마이징 가능)
  • 나머지: 발견사항·조치, 종합의견, 차회 점검, 첨부

예시값은 단양군 UPIS 사례. 양식 자체에 영향 없도록 모든 텍스트는 [대괄호] 또는 빈칸으로.
"""
from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

OUT_DIR = Path(__file__).parent


# ── 공통 데이터 ─────────────────────────────────────────────────────────────
# 예시값은 단양군 UPIS 사례 — 시안 비교용 자리표시자.
COVER_DEFAULTS = {
    "사업명": "20○○년 ○○○ 시스템 운영 및 유지관리",
    "회차": "20○○년 ○월 정기점검 (제 회차)",
    "기관": "○○군청 / ○○시청 / ○○공사 ○○○실",
    "기관담당부서": "정보화담당관 / 정보화과 등",
    "수행사": "(주)○○○",
    "작성일자": "20○○년 ○월 ○○일",
}

# 월별 커스터마이징 가능 — 행 추가/삭제·항목 변경 자유
INSPECTION_ITEMS = [
    # (분류, 점검항목, 비고)
    ("AP 서버", "CPU·메모리 정보 확인", "ap.hw.*"),
    ("AP 서버", "어댑터 / 디스크 구성 확인", "ap.hw.adapter, ap.disk.*"),
    ("AP 서버", "CPU 사용률", "ap.perf.cpu_pct"),
    ("AP 서버", "메모리 사용률", "ap.perf.mem_pct"),
    ("AP 서버", "시스템 이벤트 로그 (에러)", "ap.log.system_err"),
    ("AP 서버", "보안 이벤트 로그 (에러)", "ap.log.security_err"),
    ("AP 서버", "네트워크 라우팅 / IP", "ap.net.*"),
    ("AP 서버", "사용자 계정 점검", "ap.security.users"),
    ("AP 서버", "LED — 시스템 / PSU / 디스크 / AP 케이블", "ap.led.* (육안 M)"),
    ("GIS", "GSS 프로세스 가동 상태", "gis.gss.running"),
    ("GIS", "GSS 로그 정리 (보존정책)", "gis.gss.log_purge"),
    ("GIS", "GWS 프로세스 가동 상태", "gis.gws.running"),
    ("GIS", "GWS 로그 정리", "gis.gws.log_purge"),
    ("GIS", "GWS store 용량", "gis.gws.store_size"),
    ("GIS", "GWS HTTP 응답", "gis.gws.http"),
    ("DB", "인스턴스 가동 / 알림 로그", "db.alert"),
    ("DB", "테이블스페이스 사용률", "db.tbs"),
    ("DB", "백업 결과 / 보존", "db.backup"),
    ("DB", "Long-running 세션 / 락", "db.session"),
    ("기타", "백업 매체 보존 / 외부 반출 이력", ""),
    ("기타", "보안 패치 / 펌웨어 업데이트 검토", ""),
    ("기타", "기타 특이사항", ""),
]

FINDINGS_ROWS = 5      # 발견사항/조치 빈 행
ATTACH_ROWS = 4        # 첨부 빈 행


# ── 공용 헬퍼 ───────────────────────────────────────────────────────────────
def set_page(doc: Document, margin_cm=2.0):
    for s in doc.sections:
        s.page_width = Cm(21.0)
        s.page_height = Cm(29.7)
        s.top_margin = Cm(margin_cm)
        s.bottom_margin = Cm(margin_cm)
        s.left_margin = Cm(margin_cm)
        s.right_margin = Cm(margin_cm)


def set_run(run, *, font="맑은 고딕", size=10, bold=False, color=None):
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)
    # east-asian font binding
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)


def add_para(doc_or_cell, text="", *, align=None, **run_kw):
    p = doc_or_cell.add_paragraph()
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(text)
    set_run(r, **run_kw)
    return p


def shade_cell(cell, hex_color: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, *, sides=("top", "bottom", "left", "right"),
                     sz=8, color="000000", val="single"):
    """sz is in 1/8 pt — 8 means 1pt, 16 means 2pt."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for side in sides:
        b = tcBorders.find(qn(f"w:{side}"))
        if b is None:
            b = OxmlElement(f"w:{side}")
            tcBorders.append(b)
        b.set(qn("w:val"), val)
        b.set(qn("w:sz"), str(sz))
        b.set(qn("w:color"), color)


def set_row_height(row, cm: float, rule="atLeast"):
    trPr = row._tr.get_or_add_trPr()
    h = OxmlElement("w:trHeight")
    h.set(qn("w:val"), str(int(cm * 567)))  # 567 twip / cm
    h.set(qn("w:hRule"), rule)
    trPr.append(h)


def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "nil")
        borders.append(b)
    tblPr.append(borders)


def set_table_borders(table, *, sz=8, color="000000", inner_sz=4):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(sz))
        b.set(qn("w:color"), color)
        borders.append(b)
    for side in ("insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(inner_sz))
        b.set(qn("w:color"), color)
        borders.append(b)
    tblPr.append(borders)


def page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


# ── 공통 콘텐츠 빌더 ────────────────────────────────────────────────────────
def build_inspection_table(doc, *, header_fill: str, header_color, font="맑은 고딕",
                            zebra: bool = False, border_sz: int = 8):
    headers = ["NO", "분류", "점검 항목", "점검 결과", "특이사항", "확인"]
    widths = [Cm(1.0), Cm(1.8), Cm(6.6), Cm(2.4), Cm(4.5), Cm(1.6)]
    t = doc.add_table(rows=1, cols=len(headers))
    set_table_borders(t, sz=border_sz, inner_sz=border_sz // 2 if border_sz > 4 else 4)

    hdr = t.rows[0]
    for i, h in enumerate(headers):
        c = hdr.cells[i]
        c.width = widths[i]
        shade_cell(c, header_fill)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = c.paragraphs[0].add_run(h)
        set_run(r, font=font, size=10, bold=True, color=header_color)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_row_height(hdr, 0.9)

    for idx, (cat, item, hint) in enumerate(INSPECTION_ITEMS, start=1):
        row = t.add_row()
        cells = row.cells
        cells[0].text = ""; cells[1].text = ""; cells[2].text = ""
        cells[3].text = ""; cells[4].text = ""; cells[5].text = ""
        vals = [str(idx), cat, item, "정상 / 경고 / 오류", "", ""]
        for i, v in enumerate(vals):
            cells[i].width = widths[i]
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cells[i].paragraphs[0].text = ""
            p = cells[i].paragraphs[0]
            if i in (0, 1, 3, 5):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(v)
            set_run(r, font=font, size=9)
            if i == 2 and hint:
                p2 = cells[i].add_paragraph()
                r2 = p2.add_run(f"  · {hint}")
                set_run(r2, font=font, size=8, color=(0x80, 0x80, 0x80))
        set_row_height(row, 0.8)
        if zebra and idx % 2 == 0:
            for c in cells:
                shade_cell(c, "F5F7FA")
    return t


def build_findings_block(doc, *, header_fill: str, header_color, font="맑은 고딕",
                          border_sz: int = 8):
    add_para(doc, "발견사항 및 조치", font=font, size=12, bold=True)
    headers = ["NO", "발견 일시", "발견 내용", "원인", "조치 내용", "조치 결과"]
    widths = [Cm(1.0), Cm(2.6), Cm(4.6), Cm(3.4), Cm(4.6), Cm(1.8)]
    t = doc.add_table(rows=1, cols=len(headers))
    set_table_borders(t, sz=border_sz, inner_sz=4)
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.width = widths[i]
        shade_cell(c, header_fill)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = c.paragraphs[0].add_run(h)
        set_run(r, font=font, size=10, bold=True, color=header_color)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_row_height(t.rows[0], 0.9)
    for i in range(FINDINGS_ROWS):
        row = t.add_row()
        for j, w in enumerate(widths):
            row.cells[j].width = w
            row.cells[j].text = ""
            row.cells[j].paragraphs[0].alignment = (
                WD_ALIGN_PARAGRAPH.CENTER if j in (0, 1, 5) else WD_ALIGN_PARAGRAPH.LEFT
            )
        set_row_height(row, 1.2)
    return t


def build_summary_block(doc, *, font="맑은 고딕", border_sz: int = 8):
    add_para(doc, "종합 의견 및 차회 점검 계획", font=font, size=12, bold=True)
    t = doc.add_table(rows=4, cols=2)
    set_table_borders(t, sz=border_sz, inner_sz=4)
    labels = ["종합 의견", "권고사항", "차회 점검 예정일", "비고"]
    for i, label in enumerate(labels):
        lc = t.rows[i].cells[0]
        lc.width = Cm(3.6)
        shade_cell(lc, "F5F5F5")
        lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        lc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = lc.paragraphs[0].add_run(label)
        set_run(r, font=font, size=10, bold=True)
        rc = t.rows[i].cells[1]
        rc.width = Cm(14.4)
        rc.text = ""
        if label == "종합 의견":
            set_row_height(t.rows[i], 3.2)
        elif label == "권고사항":
            set_row_height(t.rows[i], 2.4)
        else:
            set_row_height(t.rows[i], 1.1)


def build_attach_block(doc, *, header_fill, header_color, font="맑은 고딕", border_sz=8):
    add_para(doc, "첨부", font=font, size=12, bold=True)
    headers = ["NO", "구분", "파일명 / 설명", "비고"]
    widths = [Cm(1.0), Cm(3.0), Cm(11.0), Cm(3.0)]
    t = doc.add_table(rows=1, cols=len(headers))
    set_table_borders(t, sz=border_sz, inner_sz=4)
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.width = widths[i]
        shade_cell(c, header_fill)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = c.paragraphs[0].add_run(h)
        set_run(r, font=font, size=10, bold=True, color=header_color)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_row_height(t.rows[0], 0.9)
    for i in range(ATTACH_ROWS):
        row = t.add_row()
        for j, w in enumerate(widths):
            row.cells[j].width = w
            row.cells[j].text = ""
        set_row_height(row, 0.9)


# ─────────────────────────────────────────────────────────────────────────────
#                              시안 A — 클래식 격자형
# ─────────────────────────────────────────────────────────────────────────────
def make_classic() -> Path:
    """공공기관 전통 양식 — 굵은 외곽선, 격자 위주, 도장 영역."""
    doc = Document()
    set_page(doc, margin_cm=2.2)
    font = "맑은 고딕"
    HFILL = "E8E8E8"
    HCOLOR = (0, 0, 0)

    # ── 표지 ────────────────────────────────────────────────────────────────
    add_para(doc, "")  # spacer
    add_para(doc, "")  # spacer

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("점  검  내  역  서"), font="바탕", size=32, bold=True)

    add_para(doc, "")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("INSPECTION  REPORT"), font="Arial", size=12, color=(0x55, 0x55, 0x55))

    add_para(doc, "")
    add_para(doc, "")
    add_para(doc, "")

    # 사업/회차/기관 — 굵은 격자
    cover = doc.add_table(rows=4, cols=2)
    set_table_borders(cover, sz=16, inner_sz=12)
    rows_data = [
        ("사  업  명", COVER_DEFAULTS["사업명"]),
        ("점검 회차 / 월", COVER_DEFAULTS["회차"]),
        ("기  관  명", COVER_DEFAULTS["기관"]),
        ("작  성  일", COVER_DEFAULTS["작성일자"]),
    ]
    for i, (label, val) in enumerate(rows_data):
        lc = cover.rows[i].cells[0]
        lc.width = Cm(4.5)
        shade_cell(lc, "DDDDDD")
        lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        lc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(lc.paragraphs[0].add_run(label), font="바탕", size=13, bold=True)
        rc = cover.rows[i].cells[1]
        rc.width = Cm(12.5)
        rc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_run(rc.paragraphs[0].add_run("  " + val), font="바탕", size=12)
        set_row_height(cover.rows[i], 1.4)

    add_para(doc, "")
    add_para(doc, "")
    add_para(doc, "")

    # 확인자 / 점검자 서명란 — 도장 영역 강조
    sig = doc.add_table(rows=4, cols=4)
    set_table_borders(sig, sz=16, inner_sz=12)
    sig_widths = [Cm(2.5), Cm(6.5), Cm(2.5), Cm(5.5)]
    headers_sig = [
        ("구  분", "확  인  자  (담당자 / 공무원·고객)", "구  분", "점  검  자  (수행사)"),
        ("소속·부서", "", "소속·부서", ""),
        ("성      명", "", "성      명", ""),
        ("서      명", "(인)", "서      명", "(인)"),
    ]
    for i, row in enumerate(headers_sig):
        for j, val in enumerate(row):
            c = sig.rows[i].cells[j]
            c.width = sig_widths[j]
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if i == 0 or j % 2 == 0:
                shade_cell(c, "DDDDDD")
                c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_run(c.paragraphs[0].add_run(val), font="바탕", size=11, bold=True)
            else:
                c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT if val == "(인)" else WD_ALIGN_PARAGRAPH.LEFT
                set_run(c.paragraphs[0].add_run(val if val else " "), font="바탕", size=11)
            if i == 3:
                set_row_height(sig.rows[i], 2.4)  # 도장 공간
            elif i == 0:
                set_row_height(sig.rows[i], 1.0)
            else:
                set_row_height(sig.rows[i], 1.2)

    # 표지 끝
    page_break(doc)

    # ── 2페이지: 월별 점검사항 ─────────────────────────────────────────────
    add_para(doc, "1. 월별 점검 사항", font="바탕", size=14, bold=True)
    add_para(doc, "  ※ 행은 사이트·월별로 자유롭게 추가·삭제·수정 가능. 점검 결과 칸은 정상/경고/오류 중 택1.", font=font, size=9, color=(0x66, 0x66, 0x66))
    add_para(doc, "")
    build_inspection_table(doc, header_fill=HFILL, header_color=HCOLOR, font=font, border_sz=12)

    page_break(doc)

    # ── 3페이지: 발견사항·조치 + 종합의견 ──────────────────────────────────
    add_para(doc, "2. 발견사항 및 조치", font="바탕", size=14, bold=True)
    add_para(doc, "")
    build_findings_block(doc, header_fill=HFILL, header_color=HCOLOR, font=font, border_sz=12)

    add_para(doc, "")
    add_para(doc, "3. 종합 의견 및 차회 점검 계획", font="바탕", size=14, bold=True)
    add_para(doc, "")
    build_summary_block(doc, font=font, border_sz=12)

    add_para(doc, "")
    add_para(doc, "4. 첨부", font="바탕", size=14, bold=True)
    add_para(doc, "")
    build_attach_block(doc, header_fill=HFILL, header_color=HCOLOR, font=font, border_sz=12)

    path = OUT_DIR / "점검내역서_시안A_classic.docx"
    doc.save(path)
    return path


# ─────────────────────────────────────────────────────────────────────────────
#                              시안 B — 모던 헤더형
# ─────────────────────────────────────────────────────────────────────────────
def make_modern() -> Path:
    """진한 컬러 배너 헤더, 큰 폰트, 넉넉한 여백, 카드형 레이아웃."""
    doc = Document()
    set_page(doc, margin_cm=1.8)
    font = "맑은 고딕"
    BRAND = "1F4E78"            # 진한 청색
    BRAND_LIGHT = "DCE6F1"
    HCOLOR = (0xFF, 0xFF, 0xFF)
    BRAND_RGB = (0x1F, 0x4E, 0x78)

    # 컬러 배너
    banner = doc.add_table(rows=1, cols=1)
    bc = banner.rows[0].cells[0]
    bc.width = Cm(17.4)
    shade_cell(bc, BRAND)
    bc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    bc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(bc.paragraphs[0].add_run("점검내역서"),
            font=font, size=36, bold=True, color=(0xFF, 0xFF, 0xFF))
    set_row_height(banner.rows[0], 4.0)
    remove_table_borders(banner)

    add_para(doc, "")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("Monthly Inspection Report"),
            font="Arial", size=11, color=(0x80, 0x80, 0x80))

    add_para(doc, "")
    add_para(doc, "")

    # 메타 정보 — 카드 그리드 (2 × 2)
    meta = doc.add_table(rows=2, cols=2)
    set_table_borders(meta, sz=4, color="D9D9D9", inner_sz=4)
    meta_items = [
        ("사업명", COVER_DEFAULTS["사업명"]),
        ("점검 회차", COVER_DEFAULTS["회차"]),
        ("기관 / 발주처", COVER_DEFAULTS["기관"]),
        ("작성일", COVER_DEFAULTS["작성일자"]),
    ]
    for idx, (label, val) in enumerate(meta_items):
        r, c = divmod(idx, 2)
        cell = meta.rows[r].cells[c]
        cell.width = Cm(8.6)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        p1 = cell.paragraphs[0]
        set_run(p1.add_run(label.upper()),
                font=font, size=8, bold=True, color=(0x80, 0x80, 0x80))
        p2 = cell.add_paragraph()
        set_run(p2.add_run(val), font=font, size=14, bold=True, color=BRAND_RGB)
        set_row_height(meta.rows[r], 2.4)

    add_para(doc, "")
    add_para(doc, "")

    # 확인자 / 점검자 — 모던 2분할
    sig_label = doc.add_paragraph()
    set_run(sig_label.add_run("확인자 / 점검자"),
            font=font, size=11, bold=True, color=(0x80, 0x80, 0x80))

    sig = doc.add_table(rows=4, cols=4)
    set_table_borders(sig, sz=4, color="D9D9D9", inner_sz=4)
    sig_widths = [Cm(2.4), Cm(6.4), Cm(2.4), Cm(6.2)]
    sig_data = [
        ("구분", "확인자 (담당자·공무원/고객)", "구분", "점검자 (수행사)"),
        ("소속", "", "소속", ""),
        ("성명", "", "성명", ""),
        ("서명", "", "서명", ""),
    ]
    for i, row in enumerate(sig_data):
        for j, val in enumerate(row):
            cell = sig.rows[i].cells[j]
            cell.width = sig_widths[j]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if i == 0:
                shade_cell(cell, BRAND)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_run(cell.paragraphs[0].add_run(val),
                        font=font, size=11, bold=True, color=(0xFF, 0xFF, 0xFF))
            elif j % 2 == 0:
                shade_cell(cell, BRAND_LIGHT)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_run(cell.paragraphs[0].add_run(val), font=font, size=11, bold=True, color=BRAND_RGB)
            else:
                set_run(cell.paragraphs[0].add_run(" "), font=font, size=11)
            if i == 0:
                set_row_height(sig.rows[i], 1.0)
            elif i == 3:
                set_row_height(sig.rows[i], 2.2)
            else:
                set_row_height(sig.rows[i], 1.2)

    page_break(doc)

    # 2페이지: 월별 점검사항
    title = doc.add_paragraph()
    set_run(title.add_run("01"),
            font=font, size=20, bold=True, color=BRAND_RGB)
    set_run(title.add_run("   월별 점검 사항"),
            font=font, size=16, bold=True)
    add_para(doc, "  · 사이트/월별로 행 추가·삭제·항목 수정 가능. 결과 칸은 정상/경고/오류 택1.",
             font=font, size=9, color=(0x80, 0x80, 0x80))
    add_para(doc, "")
    build_inspection_table(doc, header_fill=BRAND, header_color=HCOLOR, font=font, zebra=True, border_sz=4)

    page_break(doc)

    # 3페이지: 발견사항·조치 + 종합의견 + 첨부
    title = doc.add_paragraph()
    set_run(title.add_run("02"), font=font, size=20, bold=True, color=BRAND_RGB)
    set_run(title.add_run("   발견사항 및 조치"), font=font, size=16, bold=True)
    add_para(doc, "")
    build_findings_block(doc, header_fill=BRAND, header_color=HCOLOR, font=font, border_sz=4)

    add_para(doc, "")
    title = doc.add_paragraph()
    set_run(title.add_run("03"), font=font, size=20, bold=True, color=BRAND_RGB)
    set_run(title.add_run("   종합 의견 및 차회 점검 계획"), font=font, size=16, bold=True)
    add_para(doc, "")
    build_summary_block(doc, font=font, border_sz=4)

    add_para(doc, "")
    title = doc.add_paragraph()
    set_run(title.add_run("04"), font=font, size=20, bold=True, color=BRAND_RGB)
    set_run(title.add_run("   첨부"), font=font, size=16, bold=True)
    add_para(doc, "")
    build_attach_block(doc, header_fill=BRAND, header_color=HCOLOR, font=font, border_sz=4)

    path = OUT_DIR / "점검내역서_시안B_modern.docx"
    doc.save(path)
    return path


# ─────────────────────────────────────────────────────────────────────────────
#                       시안 C v2 — 미니멀 모던 (피드백 반영)
# ─────────────────────────────────────────────────────────────────────────────
# 변경점 (사용자 피드백):
#   1) 점검 이력 표 추가 — 이번 달이 2월이면 1월 회차가 그 위에 누적 표시
#   2) 점검사항 표 4컬럼 (NO/분류/항목/결과) — 특이사항·확인 컬럼 제거
#   3) 특이사항을 별도 자유 기술 박스로 분리
#   4) 발견사항·조치, 종합의견·차회 점검, 첨부 섹션 모두 제거
#   5) 디자인 톤 — 다크 헤로 + 가는 액센트 라인, 큰 헤딩, 표 외곽 무선·헤더만 컬러

DARK_HEX = "0F172A"          # slate-900
DARK_RGB = (0x0F, 0x17, 0x2A)
ACCENT_HEX_V2 = "10B981"     # emerald-500
ACCENT_RGB_V2 = (0x10, 0xB9, 0x81)
MUTED_RGB = (0x6B, 0x72, 0x80)
LIGHT_BG = "F8FAFC"          # slate-50
DIV_HEX = "E2E8F0"           # slate-200

HISTORY_ROWS_DEFAULT = 6     # 이전 회차 + 이번 회차 포함 누적 행
INSPECTION_ITEMS_V2 = INSPECTION_ITEMS   # 동일 데이터, 컬럼만 축소


def _section_heading(doc, eyebrow: str, title_kr: str, font: str):
    """모던 섹션 헤더 — 작은 영문 eyebrow 위, 큰 한글 타이틀 아래, 가는 액센트 라인."""
    p = doc.add_paragraph()
    set_run(p.add_run(eyebrow.upper()),
            font="Arial", size=8, bold=True, color=ACCENT_RGB_V2)
    p = doc.add_paragraph()
    set_run(p.add_run(title_kr), font=font, size=20, bold=True, color=DARK_RGB)
    # 가는 가로 액센트 라인 (1행 1셀 미니 테이블)
    line = doc.add_table(rows=1, cols=1)
    remove_table_borders(line)
    c = line.rows[0].cells[0]
    c.width = Cm(2.4)
    shade_cell(c, ACCENT_HEX_V2)
    set_row_height(line.rows[0], 0.08, rule="exact")
    c.paragraphs[0].text = ""
    add_para(doc, "")


def _modern_inspection_table(doc, font: str):
    """4컬럼 (NO / 분류 / 점검항목 / 결과) — 외곽 무선, 헤더만 다크, zebra."""
    headers = ["NO", "분류", "점검 항목", "결과"]
    widths = [Cm(1.2), Cm(2.6), Cm(10.6), Cm(3.0)]
    t = doc.add_table(rows=1, cols=len(headers))
    # 외곽 무선, 헤더 하단·내부 가로선만
    tbl = t._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "right", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "nil")
        borders.append(b)
    for side in ("bottom", "insideH"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), DIV_HEX)
        borders.append(b)
    tblPr.append(borders)

    hdr = t.rows[0]
    for i, h in enumerate(headers):
        c = hdr.cells[i]
        c.width = widths[i]
        shade_cell(c, DARK_HEX)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if i in (0, 1, 3) else WD_ALIGN_PARAGRAPH.LEFT
        # left-align 항목 컬럼은 들여쓰기
        text = ("  " + h) if i == 2 else h
        set_run(c.paragraphs[0].add_run(text),
                font=font, size=10, bold=True, color=(0xFF, 0xFF, 0xFF))
    set_row_height(hdr, 0.95)

    for idx, (cat, item, hint) in enumerate(INSPECTION_ITEMS_V2, start=1):
        row = t.add_row()
        cells = row.cells
        vals = [str(idx), cat, item, "정상 / 경고 / 오류"]
        for i, v in enumerate(vals):
            cells[i].width = widths[i]
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in (0, 1, 3) else WD_ALIGN_PARAGRAPH.LEFT
            text = ("  " + v) if i == 2 else v
            set_run(p.add_run(text), font=font, size=10, color=DARK_RGB if i != 3 else MUTED_RGB)
            if i == 2 and hint:
                p2 = cells[i].add_paragraph()
                set_run(p2.add_run("  · " + hint),
                        font=font, size=8, color=MUTED_RGB)
        set_row_height(row, 0.82)
        if idx % 2 == 0:
            for c in cells:
                shade_cell(c, LIGHT_BG)


def _history_table(doc, font: str):
    """점검 이력 — 회차/일자/점검자/요약. 이번 달이 2월이면 1월 회차가 위에."""
    headers = ["회차", "점검 일자", "점검자", "확인자", "주요 결과 요약"]
    widths = [Cm(1.5), Cm(2.6), Cm(2.6), Cm(2.6), Cm(8.1)]
    t = doc.add_table(rows=1, cols=len(headers))
    # 헤더 하단·행 사이 얇은 가로선만, 외곽 무선
    tbl = t._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "right", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "nil")
        borders.append(b)
    for side in ("bottom", "insideH"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), DIV_HEX)
        borders.append(b)
    tblPr.append(borders)

    hdr = t.rows[0]
    for i, h in enumerate(headers):
        c = hdr.cells[i]
        c.width = widths[i]
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 4 else WD_ALIGN_PARAGRAPH.CENTER
        text = ("  " + h) if i == 4 else h
        set_run(c.paragraphs[0].add_run(text),
                font=font, size=9, bold=True, color=MUTED_RGB)
    set_row_height(hdr, 0.8)

    for i in range(HISTORY_ROWS_DEFAULT):
        row = t.add_row()
        for j, w in enumerate(widths):
            row.cells[j].width = w
            row.cells[j].text = ""
            p = row.cells[j].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 4 else WD_ALIGN_PARAGRAPH.CENTER
            # 가장 최근(이번 회차) 행을 마지막 행으로 강조 표시
            if i == HISTORY_ROWS_DEFAULT - 1:
                shade_cell(row.cells[j], LIGHT_BG)
                placeholder = ["이번 회차", "", "", "", "  ← 본 점검내역서 (작성 중)"][j]
                set_run(p.add_run(placeholder),
                        font=font, size=10, bold=(j == 0), color=ACCENT_RGB_V2 if j == 0 else MUTED_RGB)
            else:
                placeholder_map = {
                    0: f"제 {HISTORY_ROWS_DEFAULT - i - 1} 회차",
                    1: "20  -  -  ",
                    2: "",
                    3: "",
                    4: "",
                }
                set_run(p.add_run(placeholder_map[j]),
                        font=font, size=10, color=MUTED_RGB if j < 4 else DARK_RGB)
        set_row_height(row, 1.0)


def _notes_box(doc, font: str, *, lines: int = 8):
    """특이사항 자유 기술 박스 — 큰 빈 영역, 얇은 회색 외곽."""
    t = doc.add_table(rows=1, cols=1)
    set_table_borders(t, sz=4, color=DIV_HEX[2:] if DIV_HEX.startswith("0x") else DIV_HEX, inner_sz=4)
    c = t.rows[0].cells[0]
    c.width = Cm(17.4)
    set_row_height(t.rows[0], lines * 0.7)
    c.text = ""
    p = c.paragraphs[0]
    set_run(p.add_run("  여기에 자유 기술 (사이트 특이상황·교체 부품·일정 변경·기타 코멘트 등)"),
            font=font, size=9, color=(0xC0, 0xC0, 0xC0))


def make_minimal() -> Path:
    doc = Document()
    set_page(doc, margin_cm=2.0)
    font = "맑은 고딕"

    # ── 표지 ────────────────────────────────────────────────────────────────
    # 다크 헤로 블록 (페이지 상단 1/3)
    hero = doc.add_table(rows=1, cols=1)
    remove_table_borders(hero)
    hc = hero.rows[0].cells[0]
    hc.width = Cm(17.0)
    shade_cell(hc, DARK_HEX)
    hc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_row_height(hero.rows[0], 7.0)

    # eyebrow
    p = hc.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run(p.add_run("    INSPECTION  REPORT"),
            font="Arial", size=9, bold=True, color=ACCENT_RGB_V2)
    # 큰 한글 타이틀
    p = hc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run(p.add_run("    점검내역서"),
            font=font, size=44, bold=True, color=(0xFF, 0xFF, 0xFF))
    # 부제
    p = hc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run(p.add_run("    월간 운영·유지관리 점검 보고"),
            font=font, size=12, color=(0xCB, 0xD5, 0xE1))

    # 액센트 가는 라인 (헤로 하단 강조)
    line = doc.add_table(rows=1, cols=1)
    remove_table_borders(line)
    lc = line.rows[0].cells[0]
    lc.width = Cm(2.4)
    shade_cell(lc, ACCENT_HEX_V2)
    set_row_height(line.rows[0], 0.08, rule="exact")
    lc.paragraphs[0].text = ""

    add_para(doc, "")
    add_para(doc, "")

    # 메타 정보 — 라벨(작은 회색) + 값(큰 진한) 카드 그리드 2x2
    meta = doc.add_table(rows=2, cols=2)
    set_table_borders(meta, sz=4, color=DIV_HEX, inner_sz=4)
    meta_items = [
        ("PROJECT", "사업명", COVER_DEFAULTS["사업명"]),
        ("PERIOD", "점검 회차 · 월", COVER_DEFAULTS["회차"]),
        ("CLIENT", "기관 / 발주처", COVER_DEFAULTS["기관"]),
        ("ISSUED", "작성일자", COVER_DEFAULTS["작성일자"]),
    ]
    for idx, (eyebrow, label, val) in enumerate(meta_items):
        r, c = divmod(idx, 2)
        cell = meta.rows[r].cells[c]
        cell.width = Cm(8.5)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        # 셀 내부 여백을 위한 빈 paragraph 제거하고 새로 구성
        cell.text = ""
        p1 = cell.paragraphs[0]
        set_run(p1.add_run("  " + eyebrow),
                font="Arial", size=8, bold=True, color=ACCENT_RGB_V2)
        p2 = cell.add_paragraph()
        set_run(p2.add_run("  " + label),
                font=font, size=8, color=MUTED_RGB)
        p3 = cell.add_paragraph()
        set_run(p3.add_run("  " + val),
                font=font, size=13, bold=True, color=DARK_RGB)
        set_row_height(meta.rows[r], 2.8)

    add_para(doc, "")
    add_para(doc, "")

    # 확인자 / 점검자 — 미니 eyebrow + 2분할 카드 (소속/성명/서명만)
    p = doc.add_paragraph()
    set_run(p.add_run("CONFIRMED  &  INSPECTED  BY"),
            font="Arial", size=8, bold=True, color=ACCENT_RGB_V2)
    p = doc.add_paragraph()
    set_run(p.add_run("확인자 · 점검자"),
            font=font, size=12, bold=True, color=DARK_RGB)
    add_para(doc, "")

    sig = doc.add_table(rows=4, cols=4)
    set_table_borders(sig, sz=4, color=DIV_HEX, inner_sz=4)
    sig_widths = [Cm(1.8), Cm(6.6), Cm(1.8), Cm(7.0)]
    sig_data = [
        ("구분", "확인자 (담당자·공무원/고객)", "구분", "점검자 (수행사)"),
        ("소속", "", "소속", ""),
        ("성명", "", "성명", ""),
        ("서명", "", "서명", ""),
    ]
    for i, row in enumerate(sig_data):
        for j, val in enumerate(row):
            cell = sig.rows[i].cells[j]
            cell.width = sig_widths[j]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if i == 0:
                shade_cell(cell, DARK_HEX)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_run(cell.paragraphs[0].add_run(val),
                        font=font, size=10, bold=True, color=(0xFF, 0xFF, 0xFF))
            elif j % 2 == 0:
                shade_cell(cell, LIGHT_BG)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_run(cell.paragraphs[0].add_run(val),
                        font=font, size=10, bold=True, color=MUTED_RGB)
            else:
                set_run(cell.paragraphs[0].add_run(" "), font=font, size=11)
            if i == 0:
                set_row_height(sig.rows[i], 0.95)
            elif i == 3:
                set_row_height(sig.rows[i], 2.4)
            else:
                set_row_height(sig.rows[i], 1.2)

    page_break(doc)

    # ── 2페이지: 점검 이력 (Inspection History) ─────────────────────────────
    _section_heading(doc, "01 · history", "점검 이력", font)
    add_para(doc, "  · 가장 최근 회차(이번 달)가 표 맨 아래에 위치. 이전 회차의 점검 일자·점검자·요약을 누적 기록.",
             font=font, size=9, color=MUTED_RGB)
    add_para(doc, "")
    _history_table(doc, font)

    add_para(doc, "")
    add_para(doc, "")

    # ── 2페이지 하단 또는 페이지 분기: 이번 달 점검 사항 ─────────────────────
    _section_heading(doc, "02 · this month", "이번 달 점검 사항", font)
    add_para(doc, "  · 결과 칸은 정상 / 경고 / 오류 중 택1. 행은 사이트·월별로 자유롭게 추가·삭제·항목 수정 가능.",
             font=font, size=9, color=MUTED_RGB)
    add_para(doc, "")
    _modern_inspection_table(doc, font)

    page_break(doc)

    # ── 3페이지: 특이사항 (자유 기술) ────────────────────────────────────────
    _section_heading(doc, "03 · notes", "특이사항", font)
    add_para(doc, "  · 점검 항목 표로 표현되지 않는 자유 코멘트 (사이트 환경, 일정 변경, 교체 부품, 후속 조치 메모 등).",
             font=font, size=9, color=MUTED_RGB)
    add_para(doc, "")
    _notes_box(doc, font, lines=14)

    path = OUT_DIR / "점검내역서_시안C_minimal_v2.docx"
    doc.save(path)
    return path


# ─────────────────────────────────────────────────────────────────────────────
#                       시안 D — annual_report 풍 (MS Word 연례보고서 템플릿 ref)
# ─────────────────────────────────────────────────────────────────────────────
# 톤: 다크 슬레이트 + 와인레드 단일 액센트 (옵션 b — 원본의 오렌지 제거)
# 구조:
#   표지     — 다크 헤로 + "점검 내역서" + 메타
#   P2       — 점검 요약 (본문 + 미니 통계 KPI 박스)
#   P3       — GIS 엔진 로그 분석 카드 그리드 (GSS / GWS / Store)  ★ 핵심 추가
#   P4       — 이번 달 점검 사항 표 + 차회 점검 계획

ANNUAL_DARK_HEX = "2F3342"          # 다크 슬레이트 (메인)
ANNUAL_DARK_RGB = (0x2F, 0x33, 0x42)
ANNUAL_NAVY_HEX = "01023B"          # 거의 검정 네이비 (표지 헤로 배경)
ANNUAL_NAVY_RGB = (0x01, 0x02, 0x3B)
ANNUAL_WINE_HEX = "A53F52"          # 와인레드 (단일 액센트)
ANNUAL_WINE_RGB = (0xA5, 0x3F, 0x52)
ANNUAL_GREY_HEX = "44546A"          # 슬레이트 그레이 (sub-text)
ANNUAL_GREY_RGB = (0x44, 0x54, 0x6A)
ANNUAL_LIGHT_HEX = "F4F4F6"         # 라이트 카드 배경
ANNUAL_DIV_HEX = "D0CECE"           # 디바이더
ANNUAL_DIV_RGB = (0xD0, 0xCE, 0xCE)
ANNUAL_PAGE_BG_HEX = "FBEEC9"       # 페이지 배경 — 큐리오시티 크림 베이지


def _annual_eyebrow_title(doc, eyebrow: str, title_kr: str, font: str):
    """모던 섹션 헤더 — 작은 영문 eyebrow (와인레드) 위, 큰 한글 타이틀 아래, 와인레드 풀폭 라인."""
    p = doc.add_paragraph()
    _tight(p)
    set_run(p.add_run(eyebrow.upper()),
            font="Arial", size=8, bold=True, color=ANNUAL_WINE_RGB)
    p = doc.add_paragraph()
    _tight(p)
    set_run(p.add_run(title_kr), font=font, size=22, bold=True, color=ANNUAL_DARK_RGB)
    # 와인레드 풀폭 가로 라인 — 본문 폭(여백 1.5cm × 2 → A4 21cm 기준 18cm)에 맞춤
    line = doc.add_table(rows=1, cols=1)
    line.autofit = False
    remove_table_borders(line)
    c = line.rows[0].cells[0]
    c.width = Cm(18.0)
    shade_cell(c, ANNUAL_WINE_HEX)
    set_row_height(line.rows[0], 0.10, rule="exact")
    _tight(c.paragraphs[0])
    c.paragraphs[0].text = ""
    add_para(doc, "")


def _set_table_total_width(table, cm: float):
    """tblW 를 명시적 dxa 값으로 강제. python-docx 의 cell.width 만으로는
    tblW 가 auto/0 로 export 되어 표가 페이지 풀 폭을 차지하지 못함."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        return
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    twip = int(cm * 567)  # 1cm ≈ 567 twip
    tblW.set(qn("w:w"), str(twip))
    tblW.set(qn("w:type"), "dxa")
    # gridCol 도 함께 보정 (1열 표 가정)
    grid = tbl.find(qn("w:tblGrid"))
    if grid is not None:
        cols = grid.findall(qn("w:gridCol"))
        if len(cols) == 1:
            cols[0].set(qn("w:w"), str(twip))


def _set_table_cell_margins(table, *, left_dxa=60, right_dxa=60, top_dxa=20, bottom_dxa=20):
    """표 전체의 default cell margin (1/20 pt 단위, 60 = 3pt ≈ 0.106cm).
    셀 내부 텍스트가 셀 width 를 초과해 우측으로 흘러나오는 문제 완화.
    """
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        return
    # 기존 tblCellMar 가 있으면 제거 후 재추가
    existing = tblPr.find(qn("w:tblCellMar"))
    if existing is not None:
        tblPr.remove(existing)
    mar = OxmlElement("w:tblCellMar")
    for side, val in (("top", top_dxa), ("left", left_dxa),
                       ("bottom", bottom_dxa), ("right", right_dxa)):
        e = OxmlElement(f"w:{side}")
        e.set(qn("w:w"), str(val))
        e.set(qn("w:type"), "dxa")
        mar.append(e)
    tblPr.append(mar)


def _set_page_background(doc, hex_color: str):
    """문서 전체 페이지 배경색 (Word 의 page color).
    document.xml 에 <w:background/> + settings.xml 에 <w:displayBackgroundShape/> 활성화.
    """
    # 1. document 레벨 <w:background w:color="...">
    body = doc.element
    existing = body.find(qn("w:background"))
    if existing is not None:
        body.remove(existing)
    bg = OxmlElement("w:background")
    bg.set(qn("w:color"), hex_color)
    body.insert(0, bg)

    # 2. settings.xml 에 displayBackgroundShape (없으면 화면에 안 보임)
    settings = doc.settings.element
    if settings.find(qn("w:displayBackgroundShape")) is None:
        disp = OxmlElement("w:displayBackgroundShape")
        settings.append(disp)


def _tight(p, *, after_pt=0, before_pt=0, line_spacing=1.0):
    """Paragraph 의 기본 space_after=8pt 등을 0 으로 reset.
    셀 안에서 vertical-align center 가 paragraph 패딩으로 인해 상단처럼 보이는 문제 해결."""
    pf = p.paragraph_format
    pf.space_before = Pt(before_pt)
    pf.space_after = Pt(after_pt)
    pf.line_spacing = line_spacing
    return p


def _set_repeat_header_row(row):
    """tr/trPr 에 <w:tblHeader/> 를 추가해 페이지 넘김 시 헤더 행 자동 반복."""
    trPr = row._tr.get_or_add_trPr()
    h = OxmlElement("w:tblHeader")
    h.set(qn("w:val"), "true")
    trPr.append(h)


def _set_cant_split(row):
    """행이 페이지 경계에서 분할되지 않도록 고정 (한 행이 한 페이지에 통째로 들어감)."""
    trPr = row._tr.get_or_add_trPr()
    cs = OxmlElement("w:cantSplit")
    trPr.append(cs)


def _annual_meta_rows(doc, font: str, items):
    """표지 메타 정보 — N행 × 2열 (label · value) 세로 배치. 보조 설명(sub) 제거.
    items = [(label, value), ...]  또는 [(label, value, _ignored_sub), ...] 호환
    """
    label_w = Cm(3.4)
    value_w = Cm(14.0)
    t = doc.add_table(rows=len(items), cols=2)
    set_table_borders(t, sz=4, color=ANNUAL_DIV_HEX, inner_sz=4)
    is_korean_label = lambda s: any('가' <= ch <= '힣' for ch in s)
    for i, item in enumerate(items):
        label, value = item[0], item[1]  # 3-tuple 호환 — sub 무시
        row = t.rows[i]
        lc = row.cells[0]
        lc.width = label_w
        shade_cell(lc, ANNUAL_DARK_HEX)                       # 다크 슬레이트 (구분 셀과 동일)
        lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        lc.text = ""
        plc = lc.paragraphs[0]
        plc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _tight(plc)
        if is_korean_label(label):
            set_run(plc.add_run(label),
                    font=font, size=10, bold=True, color=(0xFF, 0xFF, 0xFF))
        else:
            set_run(plc.add_run(label.upper()),
                    font="Arial", size=10, bold=True, color=(0xFF, 0xFF, 0xFF))
        vc = row.cells[1]
        vc.width = value_w
        vc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        vc.text = ""
        pvc = vc.paragraphs[0]
        _tight(pvc)
        set_run(pvc.add_run("  " + value),
                font=font, size=16, bold=True, color=ANNUAL_WINE_RGB)
        set_row_height(row, 1.2)


def _annual_kpi_strip(doc, font: str, kpis):
    """표지 또는 요약 페이지 KPI 4분할 스트립 (가로).
    kpis = [(label, value, sub), ...]  — 4개 권장. label·value·sub 모두 bold.
    """
    is_korean = lambda s: any('가' <= ch <= '힣' for ch in s)
    t = doc.add_table(rows=1, cols=len(kpis))
    set_table_borders(t, sz=4, color=ANNUAL_DIV_HEX, inner_sz=4)
    col_w = Cm(18.0 / len(kpis))
    for i, (label, value, sub) in enumerate(kpis):
        c = t.rows[0].cells[i]
        c.width = col_w
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        shade_cell(c, ANNUAL_LIGHT_HEX)
        c.text = ""
        # paragraph 1 — label (bold, 한글이면 맑은 고딕 / 영문이면 Arial)
        p1 = c.paragraphs[0]
        _tight(p1, after_pt=2)
        if is_korean(label):
            set_run(p1.add_run("  " + label),
                    font=font, size=10, bold=True, color=ANNUAL_GREY_RGB)
        else:
            set_run(p1.add_run("  " + label.upper()),
                    font="Arial", size=9, bold=True, color=ANNUAL_GREY_RGB)
        # paragraph 2 — value (크고 와인레드 bold)
        p2 = c.add_paragraph()
        _tight(p2, after_pt=2)
        set_run(p2.add_run("  " + value),
                font=font, size=22, bold=True, color=ANNUAL_WINE_RGB)
        # paragraph 3 — sub (bold 추가)
        p3 = c.add_paragraph()
        _tight(p3)
        set_run(p3.add_run("  " + sub),
                font=font, size=9, bold=True, color=ANNUAL_GREY_RGB)
    set_row_height(t.rows[0], 2.6)


def _annual_log_analysis_cards(doc, font: str):
    """★ GIS 엔진 로그 분석 카드 — 세로 3행 1열 (경로 한 줄 표시)."""
    cards = [
        {
            "eyebrow": "GSS",
            "title": "Spatial Server",
            "subtitle": "GeoNURIS Spatial Server",
            "path": r"\GeoNURIS_Spatial_Server\logs",
            "rows": [
                ("프로세스 가동", "정상 / 정지"),
                ("로그 정리량 (월)", "_____ MB"),
                ("30 일 ERROR", "_____ 건"),
                ("30 일 WARN", "_____ 건"),
                ("디스크 점유", "_____ %"),
            ],
        },
        {
            "eyebrow": "GWS",
            "title": "GeoWeb Server",
            "subtitle": "GeoNURIS GeoWeb Server 64",
            "path": r"C:\Program Files\GeoNURIS_GeoWeb_Server_64\logs",
            "rows": [
                ("HTTP 응답", "200 / 5xx"),
                ("Tomcat catalina ERR", "_____ 건"),
                ("WMS p95 응답", "_____ ms"),
                ("WFS p95 응답", "_____ ms"),
                ("UWES 서블릿 에러", "_____ 건"),
                ("stdout 로그 크기", "_____ MB"),
            ],
        },
        {
            "eyebrow": "STORE",
            "title": "UWES Store",
            "subtitle": "WMS/WFS 데이터 저장소",
            "path": r"C:\Program Files\GeoNURIS_GeoWeb_Server_64\webapps\uwes\store",
            "rows": [
                ("총 용량", "_____ GB"),
                ("DEM / SLOP 제외", "보존"),
                ("기타 로그 삭제", "_____ MB"),
                ("임계치 (20GB) 위반", "Y / N"),
                ("증가 추이 (월)", "_____ MB/월"),
            ],
        },
    ]

    # placeholder 도 filled 와 동일한 세로 3행 구조로 위임
    _annual_log_analysis_cards_filled(doc, font, cards)


def _annual_summary_kpis(doc, font: str):
    """이번 회차 자동 요약 KPI — 4분할 스트립 (placeholder, payload 자동 채움 가정)."""
    _annual_kpi_strip(doc, font, kpis=[
        ("점검 항목", "_____ 건", "AP·DB·GIS 합계"),
        ("정상", "_____ 건", "result = ok"),
        ("경고 / 비정상", "_____ 건", "warn + err"),
        ("육안 점검", "_____ 건", "자동수집 불가"),
    ])


def _annual_metric_chart_placeholder(doc, font: str):
    """메트릭 추이 차트 placeholder — 실제 chart 는 SW Manager 측에서 PDF/PNG embed."""
    t = doc.add_table(rows=1, cols=1)
    set_table_borders(t, sz=4, color=ANNUAL_DIV_HEX, inner_sz=4)
    c = t.rows[0].cells[0]
    c.width = Cm(17.4)
    shade_cell(c, ANNUAL_LIGHT_HEX)
    c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    c.text = ""
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("[ 메트릭 추이 차트 — 30 일 CPU·메모리·디스크 (SW Manager 자동 삽입) ]"),
            font=font, size=10, color=ANNUAL_GREY_RGB)
    set_row_height(t.rows[0], 6.0)


def _annual_inspection_table(doc, font: str):
    """4컬럼 점검 사항 표 — 시안 C v2 와 동일 패턴, 와인레드 헤더."""
    headers = ["NO", "분류", "점검 항목", "결과"]
    widths = [Cm(1.2), Cm(2.6), Cm(10.6), Cm(3.0)]
    t = doc.add_table(rows=1, cols=len(headers))

    tbl = t._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "right", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "nil")
        borders.append(b)
    for side in ("bottom", "insideH"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), ANNUAL_DIV_HEX)
        borders.append(b)
    tblPr.append(borders)

    hdr = t.rows[0]
    for i, h in enumerate(headers):
        c = hdr.cells[i]
        c.width = widths[i]
        shade_cell(c, ANNUAL_DARK_HEX)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c.paragraphs[0]
        _tight(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in (0, 1, 3) else WD_ALIGN_PARAGRAPH.LEFT
        text = ("  " + h) if i == 2 else h
        set_run(p.add_run(text),
                font=font, size=10, bold=True, color=(0xFF, 0xFF, 0xFF))
    set_row_height(hdr, 0.95)
    _set_repeat_header_row(hdr)
    _set_cant_split(hdr)

    for idx, (cat, item, hint) in enumerate(INSPECTION_ITEMS, start=1):
        row = t.add_row()
        _set_cant_split(row)
        cells = row.cells
        vals = [str(idx), cat, item, "정상 / 경고 / 오류"]
        for i, v in enumerate(vals):
            cells[i].width = widths[i]
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            _tight(p)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in (0, 1, 3) else WD_ALIGN_PARAGRAPH.LEFT
            text = ("  " + v) if i == 2 else v
            set_run(p.add_run(text),
                    font=font, size=10,
                    color=ANNUAL_DARK_RGB if i != 3 else ANNUAL_GREY_RGB)
            if i == 2 and hint:
                p2 = cells[i].add_paragraph()
                _tight(p2)
                set_run(p2.add_run("  · " + hint),
                        font=font, size=8, color=ANNUAL_GREY_RGB)
        set_row_height(row, 0.82)
        if idx % 2 == 0:
            for c in cells:
                shade_cell(c, ANNUAL_LIGHT_HEX)


def _annual_next_plan(doc, font: str):
    """차회 점검 계획 — 3블록 (일정 / 권고사항 / 후속 조치)."""
    blocks = [
        ("NEXT DATE", "차회 점검 일정",
         "20○○년 ○월 ○○일 (월 정기점검 제 ○ 회차) · 점검자 ○○○ 수행 예정"),
        ("RECOMMEND", "권고사항",
         "(자유 기술) — 본 회차 점검에서 발견된 사항에 대한 권고 / 임계치 조정 / 부품 교체 시기 등"),
        ("FOLLOW-UP", "후속 조치",
         "(자유 기술) — 발견사항 별 조치 일정 · 책임자 · 완료 예정일"),
    ]
    for _eyebrow, title, body in blocks:
        p = doc.add_paragraph()
        _tight(p, after_pt=4)
        set_run(p.add_run(title),
                font=font, size=13, bold=True, color=ANNUAL_DARK_RGB)
        p = doc.add_paragraph()
        _tight(p, after_pt=2)
        set_run(p.add_run(body),
                font=font, size=10, color=ANNUAL_DARK_RGB)
        add_para(doc, "")


def make_annual() -> Path:
    """시안 D — annual_report 풍 (다크 슬레이트 + 와인레드 단일 액센트)."""
    doc = Document()
    set_page(doc, margin_cm=1.5)
    font = "맑은 고딕"

    # ── 표지 ────────────────────────────────────────────────────────────────
    # 다크 네이비 풀블리드 헤로 (페이지 상단 절반)
    hero = doc.add_table(rows=1, cols=1)
    remove_table_borders(hero)
    hc = hero.rows[0].cells[0]
    hc.width = Cm(18.0)
    shade_cell(hc, ANNUAL_NAVY_HEX)
    hc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_row_height(hero.rows[0], 12.0)

    # eyebrow
    p = hc.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run(p.add_run("    INSPECTION  REPORT"),
            font="Arial", size=10, bold=True, color=ANNUAL_WINE_RGB)

    add_para_to_cell = hc.add_paragraph
    # 큰 한글 타이틀
    p = add_para_to_cell()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run(p.add_run("    점검 내역서"),
            font=font, size=54, bold=True, color=(0xFF, 0xFF, 0xFF))
    # 부제
    p = add_para_to_cell()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run(p.add_run("    월간 운영·유지관리 점검 보고"),
            font=font, size=14, color=(0xCB, 0xD5, 0xE1))
    p = add_para_to_cell()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run(p.add_run("    " + COVER_DEFAULTS["사업명"]),
            font=font, size=11, color=(0xCB, 0xD5, 0xE1))
    p = add_para_to_cell()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run(p.add_run("    " + COVER_DEFAULTS["기관"] + "    |    " + COVER_DEFAULTS["수행사"]),
            font=font, size=10, color=(0xCB, 0xD5, 0xE1))

    # 와인레드 가는 라인 (헤로 하단)
    line = doc.add_table(rows=1, cols=1)
    remove_table_borders(line)
    lc = line.rows[0].cells[0]
    lc.width = Cm(3.0)
    shade_cell(lc, ANNUAL_WINE_HEX)
    set_row_height(line.rows[0], 0.10, rule="exact")
    lc.paragraphs[0].text = ""

    add_para(doc, "")

    # 메타 정보 — 4행 × 3열 (label · value · sub)
    _annual_meta_rows(doc, font, items=[
        ("PERIOD", COVER_DEFAULTS["회차"].split("(")[0].strip(), "월간 정기점검"),
        ("ISSUED", COVER_DEFAULTS["작성일자"], "보고서 작성일"),
        ("SCOPE", "AP · DB · GIS", "3 tier · 멀티 호스트"),
        ("AUTO", "QR + PWA", "자동수집 채널"),
    ])

    add_para(doc, "")

    # 확인자·점검자 (간결 2분할)
    p = doc.add_paragraph()
    set_run(p.add_run("CONFIRMED  &  INSPECTED  BY"),
            font="Arial", size=8, bold=True, color=ANNUAL_WINE_RGB)

    sig = doc.add_table(rows=4, cols=4)
    set_table_borders(sig, sz=4, color=ANNUAL_DIV_HEX, inner_sz=4)
    sig_widths = [Cm(1.8), Cm(7.2), Cm(1.8), Cm(7.2)]
    sig_data = [
        ("구분", "확인자 (담당자·공무원/고객)", "구분", "점검자 (수행사)"),
        ("소속", "", "소속", ""),
        ("성명", "", "성명", ""),
        ("서명", "", "서명", ""),
    ]
    for i, row in enumerate(sig_data):
        for j, val in enumerate(row):
            cell = sig.rows[i].cells[j]
            cell.width = sig_widths[j]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if i == 0:
                shade_cell(cell, ANNUAL_DARK_HEX)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_run(cell.paragraphs[0].add_run(val),
                        font=font, size=10, bold=True, color=(0xFF, 0xFF, 0xFF))
            elif j % 2 == 0:
                shade_cell(cell, ANNUAL_LIGHT_HEX)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_run(cell.paragraphs[0].add_run(val),
                        font=font, size=10, bold=True, color=ANNUAL_GREY_RGB)
            else:
                set_run(cell.paragraphs[0].add_run(" "), font=font, size=11)
            if i == 0:
                set_row_height(sig.rows[i], 0.95)
            elif i == 3:
                set_row_height(sig.rows[i], 2.4)
            else:
                set_row_height(sig.rows[i], 1.2)

    page_break(doc)

    # ── P2 — 점검 요약 ──────────────────────────────────────────────────────
    _annual_eyebrow_title(doc, "01 · summary", "점검 요약", font)
    add_para(doc, "  · 자동수집 결과 + 점검자 코멘트를 토대로 본 회차 점검의 핵심을 정리.",
             font=font, size=9, color=ANNUAL_GREY_RGB)
    add_para(doc, "")
    _annual_summary_kpis(doc, font)
    add_para(doc, "")
    add_para(doc,
             "  본 회차는 단양군 UPIS 운영장비 (AP·DB·GIS) 의 월간 점검입니다. "
             "자동수집 (QR 반출 → PWA 스캐너 → SW Manager) 으로 받은 메트릭과 현장 육안 점검을 "
             "종합한 결과를 본 보고서에 정리합니다.",
             font=font, size=10, color=ANNUAL_DARK_RGB)
    add_para(doc, "")
    add_para(doc,
             "  (요약 본문 — 점검자가 자유 기술. 본 회차의 핵심 발견사항·조치·차회 일정 등.)",
             font=font, size=10, color=ANNUAL_GREY_RGB)

    page_break(doc)

    # ── P3 — 메트릭 추이 (별도 페이지) ──────────────────────────────────────
    _annual_eyebrow_title(doc, "02 · metrics", "메트릭 추이 (30 일)", font)
    _annual_metric_chart_placeholder(doc, font)

    page_break(doc)

    # ── P4 — GIS 엔진 로그 분석 (★ 핵심 추가) ───────────────────────────────
    _annual_eyebrow_title(doc, "03 · gis log analysis", "GIS 엔진 로그 분석", font)
    add_para(doc, "  · GSS (Spatial Server) / GWS (GeoWeb Server) / Store 의 로그 스캔 결과 요약.",
             font=font, size=9, color=ANNUAL_GREY_RGB)
    add_para(doc, "  · 출처: 점검 에이전트 자동수집 + 점검자 육안 확인.",
             font=font, size=9, color=ANNUAL_GREY_RGB)
    add_para(doc, "")
    _annual_log_analysis_cards(doc, font)
    add_para(doc, "")
    add_para(doc, "  ※ DEM / SLOP 파일은 store 의 영구 보존 대상 — 삭제 제외. 기타 로그·임시파일만 정리 대상.",
             font=font, size=9, color=ANNUAL_WINE_RGB)
    add_para(doc, "  ※ 임계치 위반 시 발견사항 표에도 자동 기록 (warn/err 결과 코드).",
             font=font, size=9, color=ANNUAL_GREY_RGB)

    page_break(doc)

    # ── P5 — 이번 달 점검 사항 ──────────────────────────────────────────────
    _annual_eyebrow_title(doc, "04 · this month", "이번 달 점검 사항", font)
    add_para(doc, "  · 결과 칸은 정상 / 경고 / 오류 중 택1. 행은 사이트·월별로 자유롭게 추가·삭제·항목 수정 가능.",
             font=font, size=9, color=ANNUAL_GREY_RGB)
    add_para(doc, "")
    _annual_inspection_table(doc, font)

    page_break(doc)

    # ── P5 — 차회 점검 계획 ─────────────────────────────────────────────────
    _annual_eyebrow_title(doc, "05 · next round", "차회 점검 계획", font)
    add_para(doc, "")
    _annual_next_plan(doc, font)

    path = OUT_DIR / "점검내역서_시안D_annual.docx"
    doc.save(path)
    return path


# ─────────────────────────────────────────────────────────────────────────────
#                       시안 D-filled — 단양군 UPIS 2026-05 예시값 완성본
# ─────────────────────────────────────────────────────────────────────────────
# 시안 D 와 동일 양식 + 모든 자리에 단양군 UPIS 점검 예시값을 임시로 채운 본.
# sample_payload.json 의 단양군 2026-05 회차 데이터와 reference 메모리의
# GSS/GWS 경로·로그 분석 예시값을 그대로 옮긴다.

FILLED_DATA = {
    "사업명": "2026년 단양군 도시계획정보체계(UPIS) 운영 및 유지관리",
    "회차": "2026년 5월 정기점검 (제 5 회차)",
    "기관": "단양군청",
    "수행사": "(주)정도유아이티",
    "작성일자": "2026년 5월 13일",
    "점검자": "박욱진",
    "확인자_소속": "단양군청",
    "확인자_성명": "(공무원 담당자)",
}

FILLED_KPI = [
    ("점검 항목", "46 건", "AP·DB·GIS 합계"),
    ("정상", "36 건", "결과 = ok"),
    ("경고 / 비정상", "2 건", "경고 = 2 / 에러 = 0"),
    ("육안 점검", "8 건", "자동수집 불가 (M)"),
]

FILLED_GIS_CARDS = [
    {
        "eyebrow": "GSS",
        "title": "Spatial Server",
        "subtitle": "GeoNURIS Spatial Server",
        "path": r"\GeoNURIS_Spatial_Server\logs",
        "rows": [
            ("프로세스 가동", "정상"),
            ("로그 정리량 (5월)", "12.4 MB"),
            ("30 일 ERROR", "0 건"),
            ("30 일 WARN", "3 건"),
            ("디스크 점유", "18 %"),
        ],
    },
    {
        "eyebrow": "GWS",
        "title": "GeoWeb Server",
        "subtitle": "GeoNURIS GeoWeb Server 64",
        "path": r"C:\Program Files\GeoNURIS_GeoWeb_Server_64\logs",
        "rows": [
            ("HTTP 응답", "200 OK"),
            ("Tomcat catalina ERR", "2 건"),
            ("WMS p95 응답", "184 ms"),
            ("WFS p95 응답", "92 ms"),
            ("UWES 서블릿 에러", "0 건"),
            ("stdout 로그 크기", "47 MB"),
        ],
    },
    {
        "eyebrow": "STORE",
        "title": "UWES Store",
        "subtitle": "WMS/WFS 데이터 저장소",
        "path": r"C:\Program Files\GeoNURIS_GeoWeb_Server_64\webapps\uwes\store",
        "rows": [
            ("총 용량", "18.4 GB"),
            ("DEM / SLOP 제외", "보존 OK"),
            ("기타 로그 삭제", "1,240 MB"),
            ("임계치 (20 GB) 위반", "N (근접)"),
            ("증가 추이 (월)", "+0.4 GB/월"),
        ],
    },
]

# (분류, 점검항목, 결과, 특이사항)
FILLED_INSPECT_ROWS = [
    ("AP 서버", "CPU·메모리 정보 확인",                    "정상", "4 core / 16 GB"),
    ("AP 서버", "어댑터 / 디스크 구성 확인",                  "정상", "어댑터 4 up / C 60.3% · D 22.1%"),
    ("AP 서버", "CPU 사용률",                              "정상", "23.4 %"),
    ("AP 서버", "메모리 사용률",                            "정상", "41.2 %"),
    ("AP 서버", "시스템 이벤트 로그 (에러)",                  "정상", "0 건"),
    ("AP 서버", "보안 이벤트 로그 (에러)",                    "정상", "0 건"),
    ("AP 서버", "네트워크 라우팅 / IP",                      "정상", "어댑터 4 up"),
    ("AP 서버", "사용자 계정 점검",                          "정상", "변동 없음"),
    ("AP 서버", "LED — 시스템 / PSU / 디스크 / AP 케이블",   "육안", "현장 육안 확인 필요"),
    ("GIS",    "GSS 프로세스 가동 상태",                     "정상", "Running"),
    ("GIS",    "GSS 로그 정리 (보존정책)",                   "정상", "12.4 MB 정리"),
    ("GIS",    "GWS 프로세스 가동 상태",                     "정상", "Running"),
    ("GIS",    "GWS 로그 정리",                            "정상", "24 개 파일 정리"),
    ("GIS",    "GWS store 용량",                          "경고", "18.4 GB (임계치 20 GB 근접)"),
    ("GIS",    "GWS HTTP 응답",                           "정상", "HTTP 200"),
    ("DB",     "인스턴스 가동 / 알림 로그",                  "정상", "errpt 신규 0 / mirror 정상"),
    ("DB",     "테이블스페이스 사용률",                      "경고", "UPIS_DATA 78.4 %"),
    ("DB",     "백업 결과 / 보존",                         "정상", "2026-05-10 export 정상"),
    ("DB",     "Long-running 세션 / 락",                  "정상", "장기 세션 없음"),
    ("기타",    "백업 매체 보존 / 외부 반출 이력",            "정상", "월간 반출 1 회 (5/10)"),
    ("기타",    "보안 패치 / 펌웨어 업데이트 검토",            "육안", "차회 검토 예정"),
    ("기타",    "기타 특이사항",                             "—",   "GWS stdout 로테이션 미설정 (권고사항 §5 참조)"),
]

FILLED_NEXT_PLAN = [
    ("차회 점검 일정", "차회 점검 일정",
     ["담당자와 협의 후 6월에 점검자 박욱진 수행 예정"]),
    ("권고사항", "권고사항",
     [
         "1) GWS geowebservice64-stdout 로그 로테이션 미설정 → logback / log4j2 size-based rotation 적용 권고.",
         "2) UWES store 용량 18.4 GB 로 임계치 20 GB 근접 — DEM·SLOP 외 임시 파일 정리 주기 단축 검토.",
         "3) Oracle 테이블스페이스 UPIS_DATA 78.4 % — 차회 점검 전 데이터 증가 추이 모니터링 + 80 % 도달 시 확장 계획 수립.",
     ]),
    ("후속 조치", "후속 조치",
     [
         "1) GWS 로그 로테이션 적용 — 수행사 6월 중순 적용 예정.",
         "2) Store 임시파일 정리 — 차회 점검 시 일괄 처리.",
         "3) Oracle 테이블스페이스 모니터링 — 일간 자동 알림 설정 (DBA 측 협의).",
     ]),
]


def _annual_kpi_strip_filled(doc, font: str, kpis):
    """필러 버전 KPI 4분할."""
    _annual_kpi_strip(doc, font, kpis=kpis)


def _annual_log_analysis_cards_filled(doc, font: str, cards_data):
    """세로 3행 1열 카드 — 한 페이지에 카드 3개 다 들어가도록 컴팩트하게.
    각 행 cantSplit + 모든 paragraph spacing 0 reset.
    """
    t = doc.add_table(rows=len(cards_data), cols=1)
    set_table_borders(t, sz=4, color=ANNUAL_DIV_HEX, inner_sz=4)

    for ri, card in enumerate(cards_data):
        _set_cant_split(t.rows[ri])
        c = t.rows[ri].cells[0]
        c.width = Cm(17.4)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c.text = ""
        _tight(c.paragraphs[0])  # 셀의 default 빈 paragraph 도 spacing 0

        # ── 헤더 inner-table (3열) ─────────────────────────────────────
        head = c.add_table(rows=1, cols=3)
        remove_table_borders(head)
        head_widths = [Cm(5.6), Cm(4.6), Cm(7.0)]

        h0 = head.rows[0].cells[0]
        h0.width = head_widths[0]
        h0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        h0.text = ""
        p = _tight(h0.paragraphs[0])
        set_run(p.add_run(card["eyebrow"]),
                font="Arial", size=8, bold=True, color=ANNUAL_WINE_RGB)
        p = _tight(h0.add_paragraph())
        set_run(p.add_run(card["title"]),
                font=font, size=13, bold=True, color=ANNUAL_DARK_RGB)

        h1 = head.rows[0].cells[1]
        h1.width = head_widths[1]
        h1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        h1.text = ""
        p = _tight(h1.paragraphs[0])
        set_run(p.add_run(card["subtitle"]),
                font=font, size=9, color=ANNUAL_GREY_RGB)

        h2 = head.rows[0].cells[2]
        h2.width = head_widths[2]
        h2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        h2.text = ""
        p = _tight(h2.paragraphs[0])
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_run(p.add_run(card["path"]),
                font="Consolas", size=7, color=ANNUAL_GREY_RGB)

        # ── 메트릭 inner-table — 2열 mini-grid (label · value 쌍 × 2) ──
        metrics = card["rows"]
        n = len(metrics)
        ncol = 2
        nrow = (n + ncol - 1) // ncol
        mt = c.add_table(rows=nrow, cols=ncol * 2)
        for k, (lbl, val) in enumerate(metrics):
            r = k % nrow
            col_base = (k // nrow) * 2
            lc2 = mt.rows[r].cells[col_base]
            lc2.width = Cm(3.8)
            lc2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            lc2.text = ""
            p = _tight(lc2.paragraphs[0])
            set_run(p.add_run("  " + lbl),
                    font=font, size=9, color=ANNUAL_GREY_RGB)
            vc2 = mt.rows[r].cells[col_base + 1]
            vc2.width = Cm(4.9)
            vc2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            vc2.text = ""
            p = _tight(vc2.paragraphs[0])
            is_alert = ("경고" in val or "위반" in val or "근접" in val
                        or "Y " in val or val.endswith("Y"))
            color = ANNUAL_WINE_RGB if is_alert else ANNUAL_DARK_RGB
            set_run(p.add_run(val),
                    font=font, size=10, bold=True, color=color)
            _set_cant_split(mt.rows[r])
        remove_table_borders(mt)

        # nested table 사이에 자동 생성되는 빈 paragraph 들도 spacing 0
        for tail_p in c.paragraphs:
            _tight(tail_p)


def doc_tables_add(cell, rows: int, cols: int):
    """cell 안에 nested table 을 추가 (python-docx 의 Cell.add_table 호환 헬퍼)."""
    return cell.add_table(rows=rows, cols=cols)


def _annual_inspection_table_filled(doc, font: str, rows_data):
    """필러 버전 — 결과·특이사항 채움. 결과 색상: 경고=와인레드.
    widths 합계 18cm 본문폭에 안전하게 맞춤 + autofit OFF + cell padding 축소.
    """
    headers = ["NO", "분류", "점검 항목", "결과", "특이사항"]
    widths = [Cm(0.9), Cm(2.2), Cm(6.6), Cm(1.6), Cm(6.0)]  # 합 17.3cm
    t = doc.add_table(rows=1, cols=len(headers))
    t.autofit = False
    _set_table_cell_margins(t, left_dxa=60, right_dxa=60)

    tbl = t._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "right", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "nil")
        borders.append(b)
    for side in ("bottom", "insideH"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), ANNUAL_DIV_HEX)
        borders.append(b)
    tblPr.append(borders)

    hdr = t.rows[0]
    for i, h in enumerate(headers):
        c = hdr.cells[i]
        c.width = widths[i]
        shade_cell(c, ANNUAL_DARK_HEX)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c.paragraphs[0]
        _tight(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in (0, 1, 3) else WD_ALIGN_PARAGRAPH.LEFT
        text = ("  " + h) if i in (2, 4) else h
        set_run(p.add_run(text),
                font=font, size=10, bold=True, color=(0xFF, 0xFF, 0xFF))
    set_row_height(hdr, 0.85)
    _set_repeat_header_row(hdr)
    _set_cant_split(hdr)

    for idx, (cat, item, result, note) in enumerate(rows_data, start=1):
        row = t.add_row()
        _set_cant_split(row)
        cells = row.cells
        vals = [str(idx), cat, item, result, note]
        for i, v in enumerate(vals):
            cells[i].width = widths[i]
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            _tight(p)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in (0, 1, 3) else WD_ALIGN_PARAGRAPH.LEFT
            text = ("  " + v) if i in (2, 4) else v
            color = ANNUAL_DARK_RGB
            bold = False
            if i == 3:
                if v == "정상":
                    color = (0x2E, 0x7D, 0x32)
                    bold = True
                elif v == "경고":
                    color = ANNUAL_WINE_RGB
                    bold = True
                elif v == "비정상":
                    color = (0xC6, 0x28, 0x28)
                    bold = True
                else:
                    color = ANNUAL_GREY_RGB
            set_run(p.add_run(text),
                    font=font, size=10, bold=bold, color=color)
        set_row_height(row, 0.7)
        if idx % 2 == 0:
            for c in cells:
                shade_cell(c, ANNUAL_LIGHT_HEX)


def make_annual_filled() -> Path:
    """시안 D 양식 + 단양군 UPIS 2026-05 예시 데이터로 채운 완성본.
    Section 1 = 표지 (margin 0, 풀 페이지 베이지). Section 2 = P2~P6 (margin 1.5cm).
    """
    doc = Document()
    set_page(doc, margin_cm=0)  # section 1 margin 0
    # header / footer 거리도 0 으로 — 그래야 페이지 풀이 베이지로 덮임
    s0 = doc.sections[0]
    s0.header_distance = Cm(0)
    s0.footer_distance = Cm(0)
    font = "맑은 고딕"

    # ══ 표지 — 페이지 풀 사이즈 베이지 박스 (21cm × 29.7cm) ══
    cover_box = doc.add_table(rows=1, cols=1)
    cover_box.autofit = False
    remove_table_borders(cover_box)
    # cell padding 은 작게 — 안쪽 여백은 nested 위·아래 paragraph 로 직접 컨트롤
    _set_table_cell_margins(cover_box, left_dxa=850, right_dxa=850,
                            top_dxa=0, bottom_dxa=0)
    _set_table_total_width(cover_box, 21.0)
    cc = cover_box.rows[0].cells[0]
    cc.width = Cm(21.0)
    shade_cell(cc, ANNUAL_PAGE_BG_HEX)
    cc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # trHeight exact — atLeast 가 Word 에서 콘텐츠 사이즈로 축소되는 문제 회피
    set_row_height(cover_box.rows[0], 29.5, rule="exact")
    _set_cant_split(cover_box.rows[0])
    cc.text = ""
    _tight(cc.paragraphs[0])

    # ── 상단 작은 헤더 (eyebrow + 우측 사업명) — nested ──
    head_eb = cc.add_table(rows=1, cols=2)
    head_eb.autofit = False
    remove_table_borders(head_eb)
    head_eb_w = [Cm(8.5), Cm(8.5)]
    he0 = head_eb.rows[0].cells[0]
    he0.width = head_eb_w[0]
    he0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    he0.text = ""
    p = _tight(he0.paragraphs[0])
    set_run(p.add_run("INSPECTION REPORT · 점검 내역서"),
            font=font, size=9, bold=True, color=ANNUAL_WINE_RGB)
    he1 = head_eb.rows[0].cells[1]
    he1.width = head_eb_w[1]
    he1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    he1.text = ""
    p = _tight(he1.paragraphs[0])
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(p.add_run(FILLED_DATA["기관"] + "  |  " + FILLED_DATA["수행사"]),
            font=font, size=9, color=ANNUAL_GREY_RGB)
    set_row_height(head_eb.rows[0], 0.9)

    # 와인레드 풀폭 라인 — nested
    line = cc.add_table(rows=1, cols=1)
    line.autofit = False
    remove_table_borders(line)
    lc = line.rows[0].cells[0]
    lc.width = Cm(17.0)
    shade_cell(lc, ANNUAL_WINE_HEX)
    set_row_height(line.rows[0], 0.10, rule="exact")
    _tight(lc.paragraphs[0])
    lc.paragraphs[0].text = ""

    # 와인 라인 다음 spacer
    p = cc.add_paragraph()
    _tight(p, before_pt=10, after_pt=0)

    # ── 매우 큰 메인 타이틀 (와인레드 72pt) ──
    p = cc.add_paragraph()
    _tight(p, after_pt=6)
    set_run(p.add_run("점검 내역서"),
            font=font, size=72, bold=True, color=ANNUAL_WINE_RGB)

    # 부제
    p = cc.add_paragraph()
    _tight(p, after_pt=2)
    set_run(p.add_run("월간 운영·유지관리 점검 보고"),
            font=font, size=14, color=ANNUAL_GREY_RGB)
    # 사업명
    p = cc.add_paragraph()
    _tight(p, after_pt=8)
    set_run(p.add_run(FILLED_DATA["사업명"]),
            font=font, size=12, bold=True, color=ANNUAL_DARK_RGB)

    # ── 본문 개요 단락 ──
    p = cc.add_paragraph()
    _tight(p, after_pt=8, line_spacing=1.4)
    set_run(p.add_run(
        "본 보고서는 단양군청 도시계획정보체계(UPIS) 운영장비의 "
        "2026년 5월 정기점검 결과를 정리합니다. "
        "AP 서버 · DB 서버 · GIS 엔진 3 tier 의 자동수집 메트릭과 현장 육안 점검을 "
        "통합하여 운영 상태와 차회 점검 시 조치 권고사항을 종합하였습니다."),
        font=font, size=11, color=ANNUAL_DARK_RGB)

    # ── 메타 3행 2열 — nested ──
    _annual_meta_rows(cc, font, items=[
        ("점검 회차", FILLED_DATA["회차"].split("(")[0].strip()),
        ("작성일", FILLED_DATA["작성일자"]),
        ("점검 범위", "AP · DB · GIS"),
    ])

    # ── 확인자 · 점검자 서명 — nested ──
    p = cc.add_paragraph()
    _tight(p, before_pt=8, after_pt=2)
    set_run(p.add_run("확인자 · 점검자"),
            font=font, size=10, bold=True, color=ANNUAL_WINE_RGB)

    sig = cc.add_table(rows=4, cols=4)
    sig.autofit = False
    set_table_borders(sig, sz=4, color=ANNUAL_DIV_HEX, inner_sz=4)
    _set_table_cell_margins(sig, left_dxa=80, right_dxa=80)
    sig_widths = [Cm(1.7), Cm(6.8), Cm(1.7), Cm(6.8)]
    sig_data = [
        ("구분", "확인자",                      "구분", "점검자"),
        ("소속", FILLED_DATA["확인자_소속"],     "소속", FILLED_DATA["수행사"]),
        ("성명", FILLED_DATA["확인자_성명"],     "성명", FILLED_DATA["점검자"]),
        ("서명", "(인)",                       "서명", "(인)"),
    ]
    for i, row in enumerate(sig_data):
        for j, val in enumerate(row):
            cell = sig.rows[i].cells[j]
            cell.width = sig_widths[j]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = _tight(cell.paragraphs[0])
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            is_label_cell = (i == 0) or (j % 2 == 0)
            if is_label_cell:
                # 구분 / 소속 / 성명 / 서명 — 모두 다크 슬레이트 + 흰 텍스트로 통일
                shade_cell(cell, ANNUAL_DARK_HEX)
                set_run(p.add_run(val),
                        font=font, size=10, bold=True, color=(0xFF, 0xFF, 0xFF))
            else:
                set_run(p.add_run(val),
                        font=font, size=10, color=ANNUAL_DARK_RGB)
            if i == 0:
                set_row_height(sig.rows[i], 0.9)
            elif i == 3:
                set_row_height(sig.rows[i], 1.5)
            else:
                set_row_height(sig.rows[i], 0.95)

    # 외부 cover_box 안의 paragraph spacing 일괄 정리
    for tail_p in cc.paragraphs:
        _tight(tail_p)

    # ── section break (nextPage) — section 2 부터 margin 1.5cm 복원 ──
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
    new_section.page_width = Cm(21.0)
    new_section.page_height = Cm(29.7)
    new_section.top_margin = Cm(1.5)
    new_section.bottom_margin = Cm(1.5)
    new_section.left_margin = Cm(1.5)
    new_section.right_margin = Cm(1.5)

    # ── P2 — 점검 요약 ──────────────────────────────────────────────────────
    _annual_eyebrow_title(doc, "01 · summary", "점검 요약", font)
    add_para(doc, "  · 자동수집 (QR 반출 → PWA 스캐너 → SW Manager) + 현장 육안 점검 결과의 핵심 정리.",
             font=font, size=9, color=ANNUAL_GREY_RGB)
    add_para(doc, "")
    _annual_kpi_strip(doc, font, FILLED_KPI)
    add_para(doc, "")
    add_para(doc,
             "  본 회차는 단양군청 도시계획정보체계(UPIS) 운영장비의 5월 정기점검 (제 5 회차) 입니다. "
             "AP 서버 (Win Server 2012R2) · DB 서버 (AIX 6.1 + Oracle 11gR2) · GIS 엔진 (GeoNURIS Spatial Server + GeoWeb Server 64) "
             "3 tier 를 점검하였으며, 자동수집 항목 46 건 중 2 건의 경고 (GWS store 용량 임계치 근접 / Oracle UPIS_DATA 테이블스페이스 78.4%) "
             "와 8 건의 육안 점검 필요 항목 (LED·케이블·펌웨어 등) 이 식별되었습니다.",
             font=font, size=10, color=ANNUAL_DARK_RGB)
    add_para(doc, "")
    add_para(doc,
             "  핵심 발견사항은 (1) GWS geowebservice64-stdout 로그 로테이션 미설정으로 인한 비대화 우려, "
             "(2) UWES store 용량의 점진적 증가, (3) Oracle UPIS_DATA 테이블스페이스 사용률 증가 추이입니다. "
             "세 항목 모두 즉시 장애 위험은 없으나 차회 점검 전까지 조치를 권고합니다 (§5 권고사항 참조).",
             font=font, size=10, color=ANNUAL_DARK_RGB)

    page_break(doc)

    # ── P3 — 메트릭 추이 (별도 페이지) ──────────────────────────────────────
    _annual_eyebrow_title(doc, "02 · metrics", "메트릭 추이 (30 일)", font)
    p = add_para(doc, "  · 자동수집 메트릭의 30 일 추이. 임계치 대비 변화 모니터링.",
                 font=font, size=9, color=ANNUAL_GREY_RGB)
    _tight(p, after_pt=4)
    # placeholder 박스 — 실제 차트는 SW Manager 자동 삽입
    t = doc.add_table(rows=1, cols=1)
    set_table_borders(t, sz=4, color=ANNUAL_DIV_HEX, inner_sz=4)
    c = t.rows[0].cells[0]
    c.width = Cm(17.4)
    shade_cell(c, ANNUAL_LIGHT_HEX)
    c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    c.text = ""
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _tight(p)
    set_run(p.add_run("[ 메트릭 추이 차트 — CPU 18~24% · 메모리 38~42% · UPIS_DATA 76→78% (SW Manager 자동 삽입 영역) ]"),
            font=font, size=10, color=ANNUAL_GREY_RGB)
    set_row_height(t.rows[0], 12.0)

    page_break(doc)

    # ── P4 — GIS 엔진 로그 분석 (한 페이지 안에 모두 표출) ─────────────────
    _annual_eyebrow_title(doc, "03 · gis log analysis", "GIS 엔진 로그 분석", font)
    p = add_para(doc, "  · GSS (Spatial Server) / GWS (GeoWeb Server) / Store 의 로그 스캔 결과 요약.",
                 font=font, size=9, color=ANNUAL_GREY_RGB)
    _tight(p, after_pt=2)
    p = add_para(doc, "  · 출처: 점검 에이전트 자동수집 (2026-05-12 22:30 KST 스캔).",
                 font=font, size=9, color=ANNUAL_GREY_RGB)
    _tight(p, after_pt=4)

    _annual_log_analysis_cards_filled(doc, font, FILLED_GIS_CARDS)

    p = add_para(doc, "  ※ DEM / SLOP 파일은 store 의 영구 보존 대상 — 삭제 제외. 기타 로그·임시파일만 정리 대상.",
                 font=font, size=9, color=ANNUAL_WINE_RGB)
    _tight(p, after_pt=2)
    p = add_para(doc, "  ※ UWES store 18.4 GB → 임계치 20 GB 근접. 차회 점검 시 일괄 정리 권고.",
                 font=font, size=9, color=ANNUAL_WINE_RGB)
    _tight(p, after_pt=2)
    p = add_para(doc, "  ※ GWS catalina ERR 2 건 — 모두 transient WMS 요청 타임아웃 (재시도 성공). 추적 불필요.",
             font=font, size=9, color=ANNUAL_GREY_RGB)
    _tight(p, after_pt=2)

    page_break(doc)

    # ── P5 — 이번 달 점검 사항 ──────────────────────────────────────────────
    _annual_eyebrow_title(doc, "04 · this month", "이번 달 점검 사항", font)
    add_para(doc, "  · 결과 색상 — 정상(녹) / 경고(와인) / 비정상(적) / 육안(회).",
             font=font, size=9, color=ANNUAL_GREY_RGB)
    add_para(doc, "")
    _annual_inspection_table_filled(doc, font, FILLED_INSPECT_ROWS)

    page_break(doc)

    # ── P5 — 차회 점검 계획 ─────────────────────────────────────────────────
    _annual_eyebrow_title(doc, "05 · next round", "차회 점검 계획", font)
    add_para(doc, "")
    for _eyebrow, title, body in FILLED_NEXT_PLAN:
        # 제목 — 검정(다크 슬레이트) 단일 표시
        p = doc.add_paragraph()
        _tight(p, after_pt=4)
        set_run(p.add_run(title),
                font=font, size=13, bold=True, color=ANNUAL_DARK_RGB)
        # body 항목 — 1)/2)/3) 각각 별도 paragraph
        if isinstance(body, str):
            body_lines = [body]
        else:
            body_lines = list(body)
        for line in body_lines:
            p = doc.add_paragraph()
            _tight(p, after_pt=2)
            set_run(p.add_run(line),
                    font=font, size=10, color=ANNUAL_DARK_RGB)
        add_para(doc, "")

    path = OUT_DIR / "점검내역서_시안D_annual_filled.docx"
    doc.save(path)
    return path


# ═════════════════════════════════════════════════════════════════════════════
# 시안 D v5 — 원본 hwpx 누락 항목 통합 (점검대상별 점검 리스트 확장)
# 시스템 구성도 제외 / 메트릭 추이 차트 제거 / 점검결과 내역 중심
# ═════════════════════════════════════════════════════════════════════════════

# 월별 점검·장애조치 이력 — 1~5월 점검완료, 6월부터 빈칸
V5_HISTORY_ROWS = [
    ("2026", "1", "23",  "정기점검", "이상없음", ""),
    ("2026", "2", "13",  "정기점검", "이상없음", ""),
    ("2026", "3", "19",  "정기점검", "이상없음", ""),
    ("2026", "4", "16",  "정기점검", "이상없음", ""),
    ("2026", "5", "13",  "정기점검", "경고 2 건 / 육안 8 건", "권고사항 §5 참조"),
    ("2026", "6",  "",   "",         "",                "" ),
    ("2026", "7",  "",   "",         "",                "" ),
    ("2026", "8",  "",   "",         "",                "" ),
    ("2026", "9",  "",   "",         "",                "" ),
    ("2026", "10", "",   "",         "",                "" ),
    ("2026", "11", "",   "",         "",                "" ),
    ("2026", "12", "",   "",         "",                "" ),
]

# 점검대상 사양 — 원본 hwpx 그대로
V5_TARGET_DB = [
    ("제품명(M/T, S/N)",  "IBM P720"),
    ("CPU",              "POWER7 3.6GHz 4 Core"),
    ("메모리",            "32 GB"),
    ("디스크",            "300 GB SAS 10Krpm × 4 (미러링 구성)"),
    ("네트워크",          "Onboard 10/100/1000-TX × 4 Port"),
    ("전원",              "이중화 전원"),
    ("운영체제",          "IBM AIX 6.1"),
]
V5_TARGET_AP = [
    ("제품명(M/T, S/N)",  "IBM X3650 M4"),
    ("CPU",              "Intel Xeon 3.6GHz × 4 Core"),
    ("메모리",            "16 GB"),
    ("디스크",            "300 GB SAS 10Krpm × 6"),
    ("네트워크",          "10/100/1000-TX × 4 Port"),
    ("전원",              "이중화 전원"),
    ("운영체제",          "WinSvrStd 2012 SNGL OLP NL 2Proc / WinSvrExtConn 2012"),
]
V5_TARGET_SW = [
    ("DBMS",     "Oracle Standard Edition Two"),
    ("GIS 엔진",  "GeoNURIS GeoSpatial Server (GSS)"),
    ("",          "GeoNURIS GeoWeb Server (GWS)"),
]

# AP 서버 점검 결과 — 원본 5열(종류/항목/내용/기준/결과) + 메모
# (종류, 점검 항목, 점검 내용, 점검 기준, 결과, 메모)
V5_AP_HEADER = [
    ("OS 정보",    "Win Server 2012 R2"),
    ("CPU",       "Intel Xeon 3.6GHz (4 Core)"),
    ("지원 업무",  "UPIS AP 서버"),
    ("Memory",    "16 GB"),
    ("Model",     "IBM X3650 M4"),
    ("Disk",      "300 GB × 6 EA"),
]
V5_AP_CHECKS = [
    ("H/W", "시스템 LED",        "Front Panel LED 육안 확인",          "적색등 유무",      "정상", "이상 없음"),
    ("H/W", "Power Supply",     "Power Supply LED 육안 확인",        "적색등 유무",      "정상", "이중화 정상"),
    ("H/W", "Disk",             "Disk LED 육안 확인",                "적색등 유무",      "정상", "6 EA 정상"),
    ("H/W", "CPU",              "작업 관리자 / 시스템 정보",            "정상 용량 확인",    "정상", "4 Core 인식"),
    ("H/W", "Memory",           "작업 관리자 / 시스템 정보",            "정상 용량 확인",    "정상", "16 GB 인식"),
    ("H/W", "Adapter",          "LED 및 케이블 연결상태 확인",          "적색등 유무",      "정상", "어댑터 4 up"),
    ("OS",  "시스템 로그",        "eventvwr.msc — System",            "Error 유무",      "정상", "에러 0 건"),
    ("OS",  "보안 로그",          "eventvwr.msc — Security",          "Error 유무",      "정상", "에러 0 건"),
    ("OS",  "Disk 여유 공간",     "내 컴퓨터 / 디스크 속성",             "용량 확인",        "정상", "C 60.3% · D 22.1%"),
    ("OS",  "Network 점검",     "netstat -r / netstat -e",          "라우팅·에러 유무",  "정상", "에러 0 건"),
    ("OS",  "IP 설정",           "ipconfig /all",                    "설정·링크 상태",    "정상", "어댑터 4 up"),
    ("보안", "사용자 계정 확인",    "lusrmgr.msc / net user",            "특이 계정 유무",    "정상", "변동 없음"),
    ("성능", "CPU 사용량",        "관리도구 → 성능 (processor %)",       "사용량 임계 이하",  "정상", "23.4 %"),
    ("성능", "Memory 사용량",     "관리도구 → 성능 (memory %)",          "사용량 임계 이하",  "정상", "41.2 %"),
]
V5_AP_EXTRA = "추가 점검 — CPU 23.4 % · Memory 41.2 % · Disk C 60.3 % / D 22.1 %"

# DB 서버(AIX) 점검 결과 — 원본의 점검방법·항목·명령어·결과 + 메모
V5_DB_HEADER = [
    ("장 비 명",    "IBM P720"),
    ("운영체제",    "AIX 6.1"),
    ("호스트이름",  "UPIS-DB"),
    ("사용 용도",   "UPIS DB 서버"),
]
V5_DB_CHECKS = [
    ("육안 점검", "H/W LED 상태 점검",      "—",                       "정상", "이상 없음"),
    ("육안 점검", "디스크 LED 상태",         "—",                       "정상", "디스크 4 정상"),
    ("육안 점검", "FAN LED 상태",           "—",                       "정상", "이상 없음"),
    ("육안 점검", "전원 LED 상태",           "—",                       "정상", "이중화 정상"),
    ("육안 점검", "케이블 연결 상태",         "—",                       "정상", "이상 없음"),
    ("구성 점검", "CPU 점검",               "lsdev -Cc processor",     "정상", "3.6GHz 4 Core"),
    ("구성 점검", "Memory 점검",            "lsdev -Cc memory",        "정상", "32 GB"),
    ("구성 점검", "Adapter 점검",           "lsdev -Cc adapter",       "정상", "어댑터 정상"),
    ("구성 점검", "Disk 점검",              "lsdev -Cc disk; lspv",    "정상", "디스크 4 정상"),
    ("구성 점검", "Network 점검",           "netstat -a",              "정상", "에러 0 건"),
    ("성능 점검", "CPU 사용률",              "topas / nmon",            "정상", "18.6 %"),
    ("성능 점검", "Memory 사용률",           "topas / nmon",            "정상", "38.4 %"),
    ("성능 점검", "SWAP 사용률",             "lsps -a",                 "정상", "8.2 %"),
    ("성능 점검", "I/O 사용률",              "iostat 1 10",             "정상", "응답 양호"),
    ("성능 점검", "네트워크 사용률",          "topas / netstat -ni",     "정상", "에러 0 건"),
    ("DATA 점검", "디스크 사용량",            "df -gH",                  "정상", "/oradata 64 %"),
    ("DATA 점검", "I-node 사용량",          "df -h",                   "정상", "임계 이하"),
    ("DATA 점검", "미러 상태 점검",            "lsvg -l rootvg",          "정상", "rootvg 정상"),
    ("네트워크",   "Link 상태 점검",          "netstat -na",             "정상", "Link Up"),
    ("네트워크",   "Ping 상태 점검",          "ping gateway",            "정상", "0% 손실"),
    ("네트워크",   "Collisions 점검",        "netstat -ni",             "정상", "0 건"),
    ("프로세서",   "각 프로세서 점검",          "topas / nmon",            "정상", "균등 분산"),
    ("로그 점검", "시스템 로그",              "errpt",                   "정상", "신규 에러 0 건"),
    ("로그 점검", "접속 로그",                "who / lastlog",           "정상", "이상 없음"),
]
V5_DB_EXTRA = "추가 점검 — CPU 18.6 % · Memory 38.4 % | 파일시스템: / 32 % · /backup 41 % · /oracle 28 % · /oradata 64 % · /Archive 19 %"

# DBMS (Oracle) 점검 결과 — 원본 17개 항목 + 명령/SQL + 결과
V5_DBMS_HEADER = [
    ("점검 대상",        "UPIS DB"),
    ("소프트웨어명",      "Oracle DB"),
    ("운영체제",          "AIX 6.1"),
    ("설치된 서버 IP",    "107.11.103.82"),
]
V5_DBMS_CHECKS = [
    ("환경",   "호스트네임",          "# hostname",                                                "정상", "UPIS-DB"),
    ("환경",   "O/S 정보",            "# oslevel -s",                                              "정상", "6100-09 확인"),
    ("환경",   "DB 버전",             "# sqlplus / -v",                                            "정상", "Oracle 11g R2"),
    ("환경",   "SID 확인",            "# echo $ORACLE_SID",                                        "정상", "UPIS"),
    ("로그",   "Oracle Alert 로그",    "# vi alert_ORT.log",                                        "정상", "에러 0 건"),
    ("아카이브", "Archive Mode 확인",   "> archive log list;",                                       "정상", "ARCHIVELOG"),
    ("아카이브", "Redo 로그 확인",       "> select * from v$logfile;",                                "정상", "3 그룹 정상"),
    ("구성",   "Control 파일 확인",    "> select * from v$controlfile;",                            "정상", "다중화 정상"),
    ("구성",   "SGA 확인",            "> show sga;",                                                "정상", "SGA 정상"),
    ("용량",   "Tablespace 상태",     "> select status from dba_tablespaces;",                     "정상", "ONLINE 전체"),
    ("용량",   "Tablespace 용량",     "> select a.tablespace_name, used % from dba_data_files …",  "경고", "UPIS_DATA 78.4 %"),
    ("용량",   "Datafile 상태",       "> select d.status, v.status from dba_data_files d …",       "정상", "AVAILABLE"),
    ("용량",   "Datafile 용량",       "> select sum(bytes) from dba_data_files;",                  "정상", "추이 안정"),
    ("백업",   "Export 백업 결과",     "# crontab -l; ls -lh expdp_*.dmp",                          "정상", "2026-05-10 export OK"),
    ("용량",   "Home 사용량",         "# df -gP $ORACLE_HOME",                                     "정상", "62 %"),
    ("용량",   "Oradata 사용량",      "# df -gP /oradata",                                         "정상", "64 %"),
    ("용량",   "Backup 사용량",       "# df -gP /backup",                                          "정상", "41 %"),
]
V5_DBMS_EXTRA = "추가 점검 — UPIS_DATA 테이블스페이스 78.4 % (임계 80 % 근접) — 차회 점검 전 모니터링 권고"

# GIS 엔진 점검 결과 — 원본의 GSS / GWS + Desktop Pro 항목 (시안 D 의 카드와 병행)
V5_GIS_CHECKS = [
    ("GSS",         "프로세스 가동 확인",          "ps -ef | grep GSS",                                     "정상", "Running"),
    ("GSS",         "로그파일 정리 (1개월 보존)",   r"/GeoNURIS_Spatial_Server/log — 1개월 이전 파일 삭제",   "정상", "12.4 MB 정리"),
    ("GSS",         "Desktop Pro 데이터저장소 구동", "Desktop Pro → 데이터저장소 → GSS 불러오기",              "정상", "정상 응답"),
    ("GWS",         "서비스 가동 확인",             "Windows 서비스 — GeoNURIS GeoWeb Server 64bit",        "정상", "Running"),
    ("GWS · STORE", "store 로그 정리",             r"…\GeoNURIS_GeoWeb_Server_64\webapps\uwes\store — DEM/SLOP 제외 삭제", "경고", "18.4 GB (임계 20 GB 근접)"),
    ("GWS",         "관리자 페이지 응답",           r"http://웹서버IP:8880/uwes → WMS Preview(Open Layers)", "정상", "HTTP 200 / 지도 표출 OK"),
]
V5_GIS_EXTRA = "추가 점검 — DEM·SLOP 파일은 영구 보존 대상 (삭제 제외). GWS stdout 로테이션 미설정 — 권고사항 §5 참조"

# 표준시스템 (UPIS 애플리케이션) 점검 — 원본 14개 항목
V5_APP_CHECKS = [
    ("도시계획",         "조회 / 검색",            "필지·고시·조서·이력·재해취약성 상세정보 조회 검색",  "정상", "정상 응답"),
    ("도시계획",         "부동산토지공부(KRAS) 연계", "토지정보 (지목·소유구분·면적) 표출 여부",          "정상", "표출 정상"),
    ("도시계획",         "토지이용계획확인서",       "해당 필지가 맞는지 표출 여부 확인",               "정상", "정상 표출"),
    ("도시계획",         "건축물대장",             "건축물대장 표출 여부",                          "정상", "정상 표출"),
    ("통계조회",         "시스템 통계",            "도시계획 통계 데이터 표출 여부",                 "정상", "정상 표출"),
    ("전자심의",         "메뉴 활성화",            "전자심의 메뉴 존재 여부",                       "정상", "정상"),
    ("지구단위계획",      "조회 / 검색",           "계획정보·규제정보·정보관리 상세정보 조회 검색",     "정상", "정상 응답"),
    ("비정형 자료실",     "메뉴 활성화",            "비정형 자료실 메뉴 존재 여부",                   "정상", "정상"),
    ("관리자",          "사용자 관리",             "사용자 승인·관리·메뉴권한 관리 메뉴 정상 여부",     "정상", "정상"),
    ("GIS 엔진 연동",   "지도 요청",               "현황도·주제도·도시계획시설 등 표출 여부",         "정상", "표출 정상"),
    ("GIS 엔진 연동",   "필지 이동",               "주소 검색 시 필지 이동",                       "정상", "정상 이동"),
    ("GIS 엔진 연동",   "하일라이팅",              "필지 이동 시 검색 필지 하일라이팅",              "정상", "정상"),
    ("GIS 엔진 연동",   "필지 정보",               "시설정보·용도지역 등 정보 표출",                 "정상", "표출 정상"),
    ("GIS 엔진 연동",   "이력 정보",               "시설정보·용도지역 이력 정보 표출",               "정상", "표출 정상"),
]

# 마지막 서명 표 — 원본 그대로
V5_SIGNATURE = {
    "사업명":   "도시계획정보체계(UPIS) 운영장비 유지관리 용역",
    "사업기간": "2026.01.09 ~ 2027.01.07",
    "점검일":   "2026년 05월 13일",
    "보고일":   "2026년 05월 13일",
    "점검자":   "㈜정도유아이티 — 박욱진",
    "확인자":   "단양군청 — (담당자)",
}


# ── v5 빌더 함수들 ─────────────────────────────────────────────────────────
def _v5_history_table(doc, font: str, rows):
    """월별 점검·장애조치 이력 표 — 6열 × 12행."""
    headers = ["년도", "월", "일", "업무", "증상", "조치내용"]
    widths = [Cm(1.8), Cm(1.2), Cm(1.2), Cm(3.6), Cm(4.6), Cm(5.0)]
    t = doc.add_table(rows=1, cols=len(headers))
    t.autofit = False
    set_table_borders(t, sz=4, color=ANNUAL_DIV_HEX, inner_sz=4)
    _set_table_cell_margins(t, left_dxa=60, right_dxa=60)

    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.width = widths[i]
        shade_cell(c, ANNUAL_DARK_HEX)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c.paragraphs[0]
        _tight(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(p.add_run(h), font=font, size=10, bold=True, color=(0xFF, 0xFF, 0xFF))
    set_row_height(t.rows[0], 0.85)
    _set_repeat_header_row(t.rows[0])

    for idx, vals in enumerate(rows, start=1):
        row = t.add_row()
        _set_cant_split(row)
        for i, v in enumerate(vals):
            cell = row.cells[i]
            cell.width = widths[i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = ""
            p = cell.paragraphs[0]
            _tight(p)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i <= 3 else WD_ALIGN_PARAGRAPH.LEFT
            text = ("  " + v) if i >= 4 and v else v
            is_warn = ("경고" in v) or ("권고" in v)
            color = ANNUAL_WINE_RGB if is_warn else ANNUAL_DARK_RGB
            set_run(p.add_run(text), font=font, size=10, color=color)
        set_row_height(row, 0.65)
        if idx % 2 == 0:
            for c in row.cells:
                shade_cell(c, ANNUAL_LIGHT_HEX)


def _v5_target_spec_table(doc, font: str, title: str, rows):
    """점검대상 사양 표 — 좌측 라벨(다크) + 우측 값. 제목은 와인레드 작은 라벨."""
    p = doc.add_paragraph()
    _tight(p, before_pt=6, after_pt=2)
    set_run(p.add_run("▷ " + title),
            font=font, size=11, bold=True, color=ANNUAL_WINE_RGB)

    t = doc.add_table(rows=len(rows), cols=2)
    t.autofit = False
    set_table_borders(t, sz=4, color=ANNUAL_DIV_HEX, inner_sz=4)
    _set_table_cell_margins(t, left_dxa=80, right_dxa=80)
    widths = [Cm(4.4), Cm(13.0)]
    for i, (lbl, val) in enumerate(rows):
        lc = t.rows[i].cells[0]
        lc.width = widths[0]
        if lbl:
            shade_cell(lc, ANNUAL_DARK_HEX)
        lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = lc.paragraphs[0]
        _tight(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if lbl:
            set_run(p.add_run(lbl),
                    font=font, size=10, bold=True, color=(0xFF, 0xFF, 0xFF))
        rc = t.rows[i].cells[1]
        rc.width = widths[1]
        rc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = rc.paragraphs[0]
        _tight(p)
        set_run(p.add_run("  " + val), font=font, size=10, color=ANNUAL_DARK_RGB)
        _set_cant_split(t.rows[i])
        set_row_height(t.rows[i], 0.7)


def _v5_tier_header_table(doc, font: str, header_pairs):
    """tier 헤더 — (라벨, 값) 쌍을 2열 2~3행 그리드로 배치."""
    n = len(header_pairs)
    rows = (n + 1) // 2
    t = doc.add_table(rows=rows, cols=4)
    t.autofit = False
    set_table_borders(t, sz=4, color=ANNUAL_DIV_HEX, inner_sz=4)
    _set_table_cell_margins(t, left_dxa=80, right_dxa=80)
    widths = [Cm(3.0), Cm(5.7), Cm(3.0), Cm(5.7)]
    for i, (lbl, val) in enumerate(header_pairs):
        r = i // 2
        c_base = (i % 2) * 2
        lc = t.rows[r].cells[c_base]
        lc.width = widths[c_base]
        shade_cell(lc, ANNUAL_DARK_HEX)
        lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = lc.paragraphs[0]
        _tight(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(p.add_run(lbl),
                font=font, size=10, bold=True, color=(0xFF, 0xFF, 0xFF))
        rc = t.rows[r].cells[c_base + 1]
        rc.width = widths[c_base + 1]
        rc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = rc.paragraphs[0]
        _tight(p)
        set_run(p.add_run("  " + val),
                font=font, size=10, color=ANNUAL_DARK_RGB)
    for r in range(rows):
        _set_cant_split(t.rows[r])
        set_row_height(t.rows[r], 0.7)


def _v5_check_table(doc, font: str, checks, *, columns):
    """점검 결과 표 — 헤더 + 결과 색상 (정상=녹, 경고=와인, 비정상=적, 육안=회).
    columns: [(label, width_cm, align), ...]  align: 'C' or 'L'
    checks 의 각 행은 columns 와 동일 길이의 tuple. '결과' 컬럼 인덱스는 자동 식별 (label 에 '결과' 포함).
    """
    result_idx = next((i for i, c in enumerate(columns) if "결과" in c[0]), -1)
    headers = [c[0] for c in columns]
    widths = [Cm(c[1]) for c in columns]
    aligns = [c[2] for c in columns]
    t = doc.add_table(rows=1, cols=len(columns))
    t.autofit = False
    set_table_borders(t, sz=4, color=ANNUAL_DIV_HEX, inner_sz=4)
    _set_table_cell_margins(t, left_dxa=60, right_dxa=60)

    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.width = widths[i]
        shade_cell(c, ANNUAL_DARK_HEX)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c.paragraphs[0]
        _tight(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(p.add_run(h), font=font, size=10, bold=True, color=(0xFF, 0xFF, 0xFF))
    set_row_height(t.rows[0], 0.78)
    _set_repeat_header_row(t.rows[0])
    _set_cant_split(t.rows[0])

    for idx, row_vals in enumerate(checks, start=1):
        row = t.add_row()
        _set_cant_split(row)
        for i, v in enumerate(row_vals):
            cell = row.cells[i]
            cell.width = widths[i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = ""
            p = cell.paragraphs[0]
            _tight(p)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if aligns[i] == "C" else WD_ALIGN_PARAGRAPH.LEFT
            text = ("  " + str(v)) if aligns[i] == "L" else str(v)
            color = ANNUAL_DARK_RGB
            bold = False
            mono = False
            if i == result_idx:
                vs = str(v).strip()
                if vs == "정상":
                    color = (0x2E, 0x7D, 0x32); bold = True
                elif vs == "경고":
                    color = ANNUAL_WINE_RGB;    bold = True
                elif vs == "비정상":
                    color = (0xC6, 0x28, 0x28); bold = True
                elif vs == "육안":
                    color = ANNUAL_GREY_RGB
            # 명령어/SQL 컬럼은 monospace
            label = columns[i][0]
            if any(k in label for k in ("명령", "SQL", "도구", "방법")):
                mono = True
            if mono:
                set_run(p.add_run(text),
                        font="Consolas", size=9, color=ANNUAL_GREY_RGB)
            else:
                set_run(p.add_run(text),
                        font=font, size=10, bold=bold, color=color)
        set_row_height(row, 0.62)
        if idx % 2 == 0:
            for c in row.cells:
                shade_cell(c, ANNUAL_LIGHT_HEX)


def _v5_extra_note(doc, font: str, text: str):
    p = doc.add_paragraph()
    _tight(p, before_pt=4, after_pt=2)
    set_run(p.add_run("  ※ " + text),
            font=font, size=9, color=ANNUAL_GREY_RGB)


def _v5_signature_table(doc, font: str, sig):
    """마지막 서명 표 — 사업명/사업기간/점검일/보고일/점검자/확인자."""
    rows = [
        ("사 업 명",  sig["사업명"],   "점 검 일",  sig["점검일"]),
        ("사업기간",  sig["사업기간"], "보 고 일",  sig["보고일"]),
        ("점 검 자",  sig["점검자"],   "확 인 자",  sig["확인자"]),
    ]
    widths = [Cm(2.4), Cm(6.3), Cm(2.4), Cm(6.3)]
    t = doc.add_table(rows=len(rows), cols=4)
    t.autofit = False
    set_table_borders(t, sz=6, color=ANNUAL_DIV_HEX, inner_sz=4)
    _set_table_cell_margins(t, left_dxa=80, right_dxa=80)
    for r, vals in enumerate(rows):
        for i, v in enumerate(vals):
            cell = t.rows[r].cells[i]
            cell.width = widths[i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            is_label = (i % 2 == 0)
            if is_label:
                shade_cell(cell, ANNUAL_DARK_HEX)
            cell.text = ""
            p = cell.paragraphs[0]
            _tight(p)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if is_label:
                set_run(p.add_run(v),
                        font=font, size=10, bold=True, color=(0xFF, 0xFF, 0xFF))
            else:
                set_run(p.add_run(v),
                        font=font, size=11, color=ANNUAL_DARK_RGB)
        _set_cant_split(t.rows[r])
        set_row_height(t.rows[r], 1.05)

    # 서명/날인 안내
    p = doc.add_paragraph()
    _tight(p, before_pt=8, after_pt=2)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(p.add_run("(서명 / 인)"),
            font=font, size=10, color=ANNUAL_GREY_RGB)


def make_annual_filled_v5() -> Path:
    """시안 D v5 — 원본 hwpx 누락 항목 통합 (점검결과 내역 확장).
    P1 표지 → P2 summary → P3 history → P4 targets →
    P5 AP → P6 DB → P7 DBMS → P8 GIS → P9 application →
    P10 next round → P11 signature.
    """
    doc = Document()
    set_page(doc, margin_cm=0)
    s0 = doc.sections[0]
    s0.header_distance = Cm(0)
    s0.footer_distance = Cm(0)
    font = "맑은 고딕"

    # ══ P1 표지 — 시안 D 표지를 그대로 (full-page beige) ══════════════════════
    cover_box = doc.add_table(rows=1, cols=1)
    cover_box.autofit = False
    remove_table_borders(cover_box)
    _set_table_cell_margins(cover_box, left_dxa=850, right_dxa=850,
                            top_dxa=0, bottom_dxa=0)
    _set_table_total_width(cover_box, 21.0)
    cc = cover_box.rows[0].cells[0]
    cc.width = Cm(21.0)
    shade_cell(cc, ANNUAL_PAGE_BG_HEX)
    cc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_row_height(cover_box.rows[0], 29.5, rule="exact")
    _set_cant_split(cover_box.rows[0])
    cc.text = ""
    _tight(cc.paragraphs[0])

    head_eb = cc.add_table(rows=1, cols=2)
    head_eb.autofit = False
    remove_table_borders(head_eb)
    he0 = head_eb.rows[0].cells[0]
    he0.width = Cm(8.5)
    he0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    he0.text = ""
    p = _tight(he0.paragraphs[0])
    set_run(p.add_run("INSPECTION REPORT · 점검 내역서"),
            font=font, size=9, bold=True, color=ANNUAL_WINE_RGB)
    he1 = head_eb.rows[0].cells[1]
    he1.width = Cm(8.5)
    he1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    he1.text = ""
    p = _tight(he1.paragraphs[0])
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(p.add_run(FILLED_DATA["기관"] + "  |  " + FILLED_DATA["수행사"]),
            font=font, size=9, color=ANNUAL_GREY_RGB)
    set_row_height(head_eb.rows[0], 0.9)

    line = cc.add_table(rows=1, cols=1)
    line.autofit = False
    remove_table_borders(line)
    lc = line.rows[0].cells[0]
    lc.width = Cm(17.0)
    shade_cell(lc, ANNUAL_WINE_HEX)
    set_row_height(line.rows[0], 0.10, rule="exact")
    _tight(lc.paragraphs[0])
    lc.paragraphs[0].text = ""

    p = cc.add_paragraph()
    _tight(p, before_pt=10, after_pt=0)
    p = cc.add_paragraph()
    _tight(p, after_pt=6)
    set_run(p.add_run("점검 내역서"),
            font=font, size=72, bold=True, color=ANNUAL_WINE_RGB)
    p = cc.add_paragraph()
    _tight(p, after_pt=2)
    set_run(p.add_run("월간 운영·유지관리 점검 보고"),
            font=font, size=14, color=ANNUAL_GREY_RGB)
    p = cc.add_paragraph()
    _tight(p, after_pt=8)
    set_run(p.add_run(FILLED_DATA["사업명"]),
            font=font, size=12, bold=True, color=ANNUAL_DARK_RGB)

    p = cc.add_paragraph()
    _tight(p, after_pt=8, line_spacing=1.4)
    set_run(p.add_run(
        "본 보고서는 단양군청 도시계획정보체계(UPIS) 운영장비의 "
        "2026년 5월 정기점검 결과를 정리합니다. "
        "AP 서버 · DB 서버 · GIS 엔진 3 tier 및 표준시스템(UPIS 애플리케이션) 의 "
        "자동수집 메트릭과 현장 육안 점검을 통합하여 운영 상태와 차회 점검 조치 권고사항을 종합하였습니다."),
        font=font, size=11, color=ANNUAL_DARK_RGB)

    _annual_meta_rows(cc, font, items=[
        ("점검 회차", FILLED_DATA["회차"].split("(")[0].strip()),
        ("작성일",    FILLED_DATA["작성일자"]),
        ("점검 범위", "AP · DB · GIS · 표준시스템"),
    ])

    p = cc.add_paragraph()
    _tight(p, before_pt=8, after_pt=2)
    set_run(p.add_run("확인자 · 점검자"),
            font=font, size=10, bold=True, color=ANNUAL_WINE_RGB)

    sig = cc.add_table(rows=4, cols=4)
    sig.autofit = False
    set_table_borders(sig, sz=4, color=ANNUAL_DIV_HEX, inner_sz=4)
    _set_table_cell_margins(sig, left_dxa=80, right_dxa=80)
    sig_widths = [Cm(1.7), Cm(6.8), Cm(1.7), Cm(6.8)]
    sig_data = [
        ("구분", "확인자",                       "구분", "점검자"),
        ("소속", FILLED_DATA["확인자_소속"],     "소속", FILLED_DATA["수행사"]),
        ("성명", FILLED_DATA["확인자_성명"],     "성명", FILLED_DATA["점검자"]),
        ("서명", "(인)",                       "서명", "(인)"),
    ]
    for i, row in enumerate(sig_data):
        for j, val in enumerate(row):
            cell = sig.rows[i].cells[j]
            cell.width = sig_widths[j]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = _tight(cell.paragraphs[0])
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            is_label_cell = (i == 0) or (j % 2 == 0)
            if is_label_cell:
                shade_cell(cell, ANNUAL_DARK_HEX)
                set_run(p.add_run(val),
                        font=font, size=10, bold=True, color=(0xFF, 0xFF, 0xFF))
            else:
                set_run(p.add_run(val),
                        font=font, size=10, color=ANNUAL_DARK_RGB)
            if i == 0:
                set_row_height(sig.rows[i], 0.9)
            elif i == 3:
                set_row_height(sig.rows[i], 1.5)
            else:
                set_row_height(sig.rows[i], 0.95)
    for tail_p in cc.paragraphs:
        _tight(tail_p)

    # section break — section 2 부터 margin 1.5cm
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
    new_section.page_width = Cm(21.0)
    new_section.page_height = Cm(29.7)
    new_section.top_margin = Cm(1.5)
    new_section.bottom_margin = Cm(1.5)
    new_section.left_margin = Cm(1.5)
    new_section.right_margin = Cm(1.5)

    # ══ P2 — 점검 요약 (시안 D 의 P2) ════════════════════════════════════════
    _annual_eyebrow_title(doc, "01 · summary", "점검 요약", font)
    add_para(doc, "  · 자동수집 (QR 반출 → PWA 스캐너 → SW Manager) + 현장 육안 점검 결과의 핵심 정리.",
             font=font, size=9, color=ANNUAL_GREY_RGB)
    add_para(doc, "")
    _annual_kpi_strip(doc, font, FILLED_KPI)
    add_para(doc, "")
    add_para(doc,
             "  본 회차는 단양군청 도시계획정보체계(UPIS) 운영장비의 5월 정기점검 (제 5 회차) 입니다. "
             "AP 서버 (Win Server 2012 R2) · DB 서버 (AIX 6.1 + Oracle 11g R2) · GIS 엔진 (GeoNURIS GSS + GWS 64) "
             "3 tier 와 표준시스템(UPIS 애플리케이션) 14 개 메뉴를 점검하였으며, 자동수집 46 건 중 2 건의 경고 "
             "(GWS store 18.4 GB → 임계 20 GB 근접 / Oracle UPIS_DATA 78.4 %) 와 8 건의 육안 점검 필요 항목이 식별되었습니다.",
             font=font, size=10, color=ANNUAL_DARK_RGB)
    add_para(doc, "")
    add_para(doc,
             "  핵심 발견사항은 (1) GWS geowebservice64-stdout 로그 로테이션 미설정, "
             "(2) UWES store 용량의 점진 증가, (3) Oracle UPIS_DATA 테이블스페이스 사용률 증가 추이입니다. "
             "세 항목 모두 즉시 장애 위험은 없으나 차회 점검 전까지 조치를 권고합니다 (§9 차회 점검 계획 참조).",
             font=font, size=10, color=ANNUAL_DARK_RGB)
    page_break(doc)

    # ══ P3 — 월별 점검·장애조치 이력 요약 ═════════════════════════════════════
    _annual_eyebrow_title(doc, "02 · history", "월별 점검·장애조치 이력 요약", font)
    add_para(doc, "  · 2026년 연간 정기점검 및 장애조치 이력 — 1~5월 점검 완료, 6월 이후 예정.",
             font=font, size=9, color=ANNUAL_GREY_RGB)
    add_para(doc, "")
    _v5_history_table(doc, font, V5_HISTORY_ROWS)
    page_break(doc)

    # ══ P4 — 점검대상 (사양 명세) ═════════════════════════════════════════════
    _annual_eyebrow_title(doc, "03 · targets", "점검대상", font)
    add_para(doc, "  · DB 서버 / AP 서버 / 소프트웨어 — 본 회차 점검대상 H/W·S/W 사양.",
             font=font, size=9, color=ANNUAL_GREY_RGB)
    _v5_target_spec_table(doc, font, "DB 서버", V5_TARGET_DB)
    _v5_target_spec_table(doc, font, "AP 서버", V5_TARGET_AP)
    _v5_target_spec_table(doc, font, "소프트웨어", V5_TARGET_SW)
    page_break(doc)

    # ══ P5 — AP 서버 점검 결과 ════════════════════════════════════════════════
    _annual_eyebrow_title(doc, "04 · ap server", "AP 서버 점검 결과", font)
    add_para(doc, "  · IBM X3650 M4 / Win Server 2012 R2 / UPIS AP 서버.",
             font=font, size=9, color=ANNUAL_GREY_RGB)
    add_para(doc, "")
    _v5_tier_header_table(doc, font, V5_AP_HEADER)
    add_para(doc, "")
    _v5_check_table(doc, font, V5_AP_CHECKS, columns=[
        ("종류",       1.4, "C"),
        ("점검 항목",   3.4, "L"),
        ("점검 방법",   5.4, "L"),
        ("점검 기준",   3.6, "L"),
        ("결과",        1.2, "C"),
        ("메모",        2.4, "L"),
    ])
    _v5_extra_note(doc, font, V5_AP_EXTRA)
    page_break(doc)

    # ══ P6 — DB 서버(AIX) 점검 결과 ═══════════════════════════════════════════
    _annual_eyebrow_title(doc, "05 · db server", "DB 서버 (AIX) 점검 결과", font)
    add_para(doc, "  · IBM P720 / AIX 6.1 / UPIS DB 서버 — 육안·구성·성능·DATA·네트워크·프로세서·로그.",
             font=font, size=9, color=ANNUAL_GREY_RGB)
    add_para(doc, "")
    _v5_tier_header_table(doc, font, V5_DB_HEADER)
    add_para(doc, "")
    _v5_check_table(doc, font, V5_DB_CHECKS, columns=[
        ("점검 방법",   1.8, "C"),
        ("점검 항목",   4.0, "L"),
        ("명령어",      6.0, "L"),
        ("결과",        1.2, "C"),
        ("메모",        4.4, "L"),
    ])
    _v5_extra_note(doc, font, V5_DB_EXTRA)
    page_break(doc)

    # ══ P7 — DBMS (Oracle) 점검 결과 ══════════════════════════════════════════
    _annual_eyebrow_title(doc, "06 · dbms", "DBMS (Oracle) 점검 결과", font)
    add_para(doc, "  · Oracle Standard Edition Two / 11g R2 — 환경·로그·아카이브·구성·용량·백업.",
             font=font, size=9, color=ANNUAL_GREY_RGB)
    add_para(doc, "")
    _v5_tier_header_table(doc, font, V5_DBMS_HEADER)
    add_para(doc, "")
    _v5_check_table(doc, font, V5_DBMS_CHECKS, columns=[
        ("구분",       1.4, "C"),
        ("점검 항목",   3.6, "L"),
        ("명령 / SQL", 7.2, "L"),
        ("결과",       1.2, "C"),
        ("메모",       4.0, "L"),
    ])
    _v5_extra_note(doc, font, V5_DBMS_EXTRA)
    page_break(doc)

    # ══ P8 — GIS 엔진 점검 결과 ═══════════════════════════════════════════════
    _annual_eyebrow_title(doc, "07 · gis engine", "GIS 엔진 점검 결과", font)
    add_para(doc, "  · GeoNURIS GSS (Spatial Server) / GWS (GeoWeb Server 64) / Desktop Pro 연계.",
             font=font, size=9, color=ANNUAL_GREY_RGB)
    add_para(doc, "")
    _v5_check_table(doc, font, V5_GIS_CHECKS, columns=[
        ("대상",       2.0, "C"),
        ("점검 항목",   3.6, "L"),
        ("점검 방법",   7.4, "L"),
        ("결과",       1.2, "C"),
        ("메모",       3.2, "L"),
    ])
    _v5_extra_note(doc, font, V5_GIS_EXTRA)
    add_para(doc, "")
    # UWES Store 카드 — alert 핵심 메시지 (임계 20 GB 근접) 시각적 강조용
    p = add_para(doc, "  · 보조 — UWES Store 자동수집 메트릭 (30 일 요약).",
                 font=font, size=9, color=ANNUAL_GREY_RGB)
    _tight(p, after_pt=4)
    store_only = [c for c in FILLED_GIS_CARDS if c["eyebrow"] == "STORE"]
    _annual_log_analysis_cards_filled(doc, font, store_only)
    page_break(doc)

    # ══ P9 — 표준시스템 (UPIS 애플리케이션) 점검 결과 ═════════════════════════
    _annual_eyebrow_title(doc, "08 · application", "표준시스템 (UPIS 애플리케이션) 점검 결과", font)
    add_para(doc, "  · 도시계획 / 통계 / 전자심의 / 지구단위계획 / 비정형 / 관리자 / GIS 연동 — 14 개 메뉴.",
             font=font, size=9, color=ANNUAL_GREY_RGB)
    add_para(doc, "")
    _v5_check_table(doc, font, V5_APP_CHECKS, columns=[
        ("분류",         3.0, "C"),
        ("세부 분류",     3.6, "L"),
        ("점검 내용",     7.6, "L"),
        ("결과",         1.2, "C"),
        ("메모",         2.0, "L"),
    ])
    page_break(doc)

    # ══ P10 — 차회 점검 계획 (시안 D 의 next round) ═══════════════════════════
    _annual_eyebrow_title(doc, "09 · next round", "차회 점검 계획", font)
    add_para(doc, "")
    for _eyebrow, title, body in FILLED_NEXT_PLAN:
        p = doc.add_paragraph()
        _tight(p, after_pt=4)
        set_run(p.add_run(title),
                font=font, size=13, bold=True, color=ANNUAL_DARK_RGB)
        body_lines = [body] if isinstance(body, str) else list(body)
        for line in body_lines:
            p = doc.add_paragraph()
            _tight(p, after_pt=2)
            set_run(p.add_run(line),
                    font=font, size=10, color=ANNUAL_DARK_RGB)
        add_para(doc, "")
    page_break(doc)

    # ══ P11 — 서명 (사업명/사업기간/점검일/보고일/점검자/확인자) ══════════════
    _annual_eyebrow_title(doc, "10 · signature", "확인 · 서명", font)
    add_para(doc, "  · 본 점검내역서의 사업 정보 및 점검·확인자 서명.",
             font=font, size=9, color=ANNUAL_GREY_RGB)
    add_para(doc, "")
    _v5_signature_table(doc, font, V5_SIGNATURE)

    path = OUT_DIR / "점검내역서_시안D_v5.docx"
    doc.save(path)
    return path


# ─────────────────────────────────────────────────────────────────────────────
# 시안 D v6 — inspection-report-d-v6 (2026-05-17)
#   변경점: P5 메트릭 추이 신규, P9 점검표 / P10 카드 페이지 분할, P11 서명 제거
#   v5 의 헬퍼·데이터 상수 (FILLED_*, V5_*, _v5_*) 와 v4 의 _annual_log_analysis_cards_filled 재사용
# ─────────────────────────────────────────────────────────────────────────────

V6_METRICS_KPI = [
    ("CPU 평균", "22 %", "30일 평균"),
    ("MEM 평균", "40 %", "30일 평균"),
    ("DISK 최대", "78 %", "30일 worst"),
    ("수집 일수", "28 / 30", "agent snapshot"),
]

V6_GIS_CARDS = [
    {
        "eyebrow": "GSS",
        "title": "Spatial Server",
        "subtitle": "GeoNURIS Spatial Server",
        "path": r"\GeoNURIS_Spatial_Server\logs",
        "rows": [
            ("프로세스 가동",  "정상"),
            ("로그 정리량 (월)", "120 MB"),
            ("30 일 ERROR",   "0 건"),
            ("30 일 WARN",    "3 건"),
            ("디스크 점유",    "42 %"),
        ],
    },
    {
        "eyebrow": "GWS",
        "title": "GeoWeb Server",
        "subtitle": "GeoNURIS GeoWeb Server 64",
        "path": r"C:\Program Files\GeoNURIS_GeoWeb_Server_64\logs",
        "rows": [
            ("HTTP 응답",          "200 OK"),
            ("Tomcat catalina ERR", "2 건"),
            ("stdout 로그 크기",    "320 MB"),
            ("UWES DEM/SLOP",      "보존"),
        ],
    },
    {
        "eyebrow": "STORE",
        "title": "UWES Store",
        "subtitle": "WMS/WFS 데이터 저장소",
        "path": r"...\webapps\uwes\store",
        "rows": [
            ("총 용량",              "18.4 GB (임계근접)"),
            ("DEM / SLOP 제외",      "보존"),
            ("기타 로그 삭제",       "240 MB"),
            ("임계치 (20GB) 위반",   "근접"),
            ("증가 추이 (월)",       "1.2 GB/월"),
        ],
    },
]


def _v6_metric_section(doc, font: str, kpis):
    """P5 신규 — 30일 메트릭 추이. KPI 4분할 + 차트 placeholder(8.0cm, 디자인팀 R6) + 수집 대기 캡션."""
    _annual_eyebrow_title(doc, "04 · metrics", "메트릭 추이 (30일 CPU·메모리·디스크)", font)
    add_para(doc, "  · agent 수집 누적 — 호스트별 line 분리, 와인/슬레이트/다크 3색 라인.",
             font=font, size=9, color=ANNUAL_GREY_RGB)
    add_para(doc, "")
    _annual_kpi_strip(doc, font, kpis)
    add_para(doc, "")

    # 차트 placeholder 박스 — 17.4 × 8.0 cm (디자인팀 자문 R6: v4 12cm → -4cm)
    t = doc.add_table(rows=1, cols=1)
    set_table_borders(t, sz=4, color=ANNUAL_DIV_HEX, inner_sz=4)
    c = t.rows[0].cells[0]
    c.width = Cm(17.4)
    shade_cell(c, ANNUAL_LIGHT_HEX)
    c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    c.text = ""
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("[ 30일 CPU·메모리·디스크 라인 차트 — SW Manager 자동 삽입 (17.4 × 8.0 cm) ]"),
            font=font, size=10, color=ANNUAL_GREY_RGB)
    set_row_height(t.rows[0], 8.0)

    # 수집 대기 캡션 (디자인팀 R3 — italic, 향후 PWA UI 에서 dashed border-left 추가)
    add_para(doc, "")
    p = doc.add_paragraph()
    _tight(p, before_pt=2, after_pt=0)
    run = p.add_run("  ※ 수집 대기 (2 일) — agent 미회신 구간은 차트에 공백으로 표시됨.")
    run.italic = True
    set_run(run, font=font, size=9, color=ANNUAL_GREY_RGB)


def make_annual_filled_v6() -> Path:
    """시안 D v6 — v5 기반 + P5 메트릭 신규 + P10 GIS 카드 분할 + P13 서명 제거.

    P1 표지 → P2 summary → P3 history → P4 targets →
    P5 metrics (신규) → P6 AP → P7 DB → P8 DBMS →
    P9 GIS 점검표 → P10 GIS 카드 (신규) → P11 application → P12 next round.
    """
    doc = Document()
    set_page(doc, margin_cm=0)
    s0 = doc.sections[0]
    s0.header_distance = Cm(0)
    s0.footer_distance = Cm(0)
    font = "맑은 고딕"

    # ══ P1 표지 (v5 그대로) ════════════════════════════════════════════════════
    cover_box = doc.add_table(rows=1, cols=1)
    cover_box.autofit = False
    remove_table_borders(cover_box)
    _set_table_cell_margins(cover_box, left_dxa=850, right_dxa=850,
                            top_dxa=0, bottom_dxa=0)
    _set_table_total_width(cover_box, 21.0)
    cc = cover_box.rows[0].cells[0]
    cc.width = Cm(21.0)
    shade_cell(cc, ANNUAL_PAGE_BG_HEX)
    cc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_row_height(cover_box.rows[0], 29.5, rule="exact")
    _set_cant_split(cover_box.rows[0])
    cc.text = ""
    _tight(cc.paragraphs[0])

    head_eb = cc.add_table(rows=1, cols=2)
    head_eb.autofit = False
    remove_table_borders(head_eb)
    he0 = head_eb.rows[0].cells[0]
    he0.width = Cm(8.5)
    he0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    he0.text = ""
    p = _tight(he0.paragraphs[0])
    set_run(p.add_run("INSPECTION REPORT · 점검 내역서"),
            font=font, size=9, bold=True, color=ANNUAL_WINE_RGB)
    he1 = head_eb.rows[0].cells[1]
    he1.width = Cm(8.5)
    he1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    he1.text = ""
    p = _tight(he1.paragraphs[0])
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(p.add_run(FILLED_DATA["기관"] + "  |  " + FILLED_DATA["수행사"]),
            font=font, size=9, color=ANNUAL_GREY_RGB)
    set_row_height(head_eb.rows[0], 0.9)

    line = cc.add_table(rows=1, cols=1)
    line.autofit = False
    remove_table_borders(line)
    lc = line.rows[0].cells[0]
    lc.width = Cm(17.0)
    shade_cell(lc, ANNUAL_WINE_HEX)
    set_row_height(line.rows[0], 0.10, rule="exact")
    _tight(lc.paragraphs[0])
    lc.paragraphs[0].text = ""

    p = cc.add_paragraph()
    _tight(p, before_pt=10, after_pt=0)
    p = cc.add_paragraph()
    _tight(p, after_pt=6)
    set_run(p.add_run("점검 내역서"),
            font=font, size=72, bold=True, color=ANNUAL_WINE_RGB)
    p = cc.add_paragraph()
    _tight(p, after_pt=2)
    set_run(p.add_run("월간 운영·유지관리 점검 보고"),
            font=font, size=14, color=ANNUAL_GREY_RGB)
    p = cc.add_paragraph()
    _tight(p, after_pt=8)
    set_run(p.add_run(FILLED_DATA["사업명"]),
            font=font, size=12, bold=True, color=ANNUAL_DARK_RGB)

    p = cc.add_paragraph()
    _tight(p, after_pt=8, line_spacing=1.4)
    set_run(p.add_run(
        "본 보고서는 단양군청 도시계획정보체계(UPIS) 운영장비의 "
        "2026년 5월 정기점검 결과를 정리합니다. "
        "AP 서버 · DB 서버 · GIS 엔진 3 tier 및 표준시스템(UPIS 애플리케이션) 의 "
        "자동수집 메트릭과 현장 육안 점검을 통합하여 운영 상태와 차회 점검 조치 권고사항을 종합하였습니다."),
        font=font, size=11, color=ANNUAL_DARK_RGB)

    _annual_meta_rows(cc, font, items=[
        ("점검 회차", FILLED_DATA["회차"].split("(")[0].strip()),
        ("작성일",    FILLED_DATA["작성일자"]),
        ("점검 범위", "AP · DB · GIS · 표준시스템"),
    ])

    p = cc.add_paragraph()
    _tight(p, before_pt=8, after_pt=2)
    set_run(p.add_run("확인자 · 점검자"),
            font=font, size=10, bold=True, color=ANNUAL_WINE_RGB)

    sig = cc.add_table(rows=4, cols=4)
    sig.autofit = False
    set_table_borders(sig, sz=4, color=ANNUAL_DIV_HEX, inner_sz=4)
    _set_table_cell_margins(sig, left_dxa=80, right_dxa=80)
    sig_widths = [Cm(1.7), Cm(6.8), Cm(1.7), Cm(6.8)]
    sig_data = [
        ("구분", "확인자",                   "구분", "점검자"),
        ("소속", FILLED_DATA["확인자_소속"], "소속", FILLED_DATA["수행사"]),
        ("성명", FILLED_DATA["확인자_성명"], "성명", FILLED_DATA["점검자"]),
        ("서명", "(인)",                     "서명", "(인)"),
    ]
    for i, row in enumerate(sig_data):
        for j, val in enumerate(row):
            cell = sig.rows[i].cells[j]
            cell.width = sig_widths[j]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = _tight(cell.paragraphs[0])
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            is_label_cell = (i == 0) or (j % 2 == 0)
            if is_label_cell:
                shade_cell(cell, ANNUAL_DARK_HEX)
                set_run(p.add_run(val),
                        font=font, size=10, bold=True, color=(0xFF, 0xFF, 0xFF))
            else:
                set_run(p.add_run(val),
                        font=font, size=10, color=ANNUAL_DARK_RGB)
            if i == 0:    set_row_height(sig.rows[i], 0.9)
            elif i == 3:  set_row_height(sig.rows[i], 1.5)
            else:         set_row_height(sig.rows[i], 0.95)
    for tail_p in cc.paragraphs:
        _tight(tail_p)

    # section break
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
    new_section.page_width = Cm(21.0)
    new_section.page_height = Cm(29.7)
    new_section.top_margin = Cm(1.5)
    new_section.bottom_margin = Cm(1.5)
    new_section.left_margin = Cm(1.5)
    new_section.right_margin = Cm(1.5)

    # ══ P2 summary ════════════════════════════════════════════════════════════
    _annual_eyebrow_title(doc, "01 · summary", "점검 요약", font)
    add_para(doc, "  · 자동수집 (QR 반출 → PWA 스캐너 → SW Manager) + 현장 육안 점검 결과의 핵심 정리.",
             font=font, size=9, color=ANNUAL_GREY_RGB)
    add_para(doc, "")
    _annual_kpi_strip(doc, font, FILLED_KPI)
    add_para(doc, "")
    add_para(doc,
             "  본 회차는 단양군청 도시계획정보체계(UPIS) 운영장비의 5월 정기점검 (제 5 회차) 입니다. "
             "AP 서버 (Win Server 2012 R2) · DB 서버 (AIX 6.1 + Oracle 11g R2) · GIS 엔진 (GeoNURIS GSS + GWS 64) "
             "3 tier 와 표준시스템(UPIS 애플리케이션) 14 개 메뉴를 점검하였으며, 자동수집 46 건 중 2 건의 경고 "
             "(GWS store 18.4 GB → 임계 20 GB 근접 / Oracle UPIS_DATA 78.4 %) 와 8 건의 육안 점검 필요 항목이 식별되었습니다.",
             font=font, size=10, color=ANNUAL_DARK_RGB)
    add_para(doc, "")
    add_para(doc,
             "  핵심 발견사항은 (1) GWS geowebservice64-stdout 로그 로테이션 미설정, "
             "(2) UWES store 용량의 점진 증가, (3) Oracle UPIS_DATA 테이블스페이스 사용률 증가 추이입니다. "
             "세 항목 모두 즉시 장애 위험은 없으나 차회 점검 전까지 조치를 권고합니다 (§10 차회 점검 계획 참조).",
             font=font, size=10, color=ANNUAL_DARK_RGB)
    page_break(doc)

    # ══ P3 history ════════════════════════════════════════════════════════════
    _annual_eyebrow_title(doc, "02 · history", "월별 점검·장애조치 이력 요약", font)
    add_para(doc, "  · 2026년 연간 정기점검 및 장애조치 이력 — 1~5월 점검 완료, 6월 이후 예정.",
             font=font, size=9, color=ANNUAL_GREY_RGB)
    add_para(doc, "")
    _v5_history_table(doc, font, V5_HISTORY_ROWS)
    page_break(doc)

    # ══ P4 targets ════════════════════════════════════════════════════════════
    _annual_eyebrow_title(doc, "03 · targets", "점검대상", font)
    add_para(doc, "  · DB 서버 / AP 서버 / 소프트웨어 — 본 회차 점검대상 H/W·S/W 사양.",
             font=font, size=9, color=ANNUAL_GREY_RGB)
    _v5_target_spec_table(doc, font, "DB 서버", V5_TARGET_DB)
    _v5_target_spec_table(doc, font, "AP 서버", V5_TARGET_AP)
    _v5_target_spec_table(doc, font, "소프트웨어", V5_TARGET_SW)
    page_break(doc)

    # ══ P5 metrics (v6 신규) ══════════════════════════════════════════════════
    _v6_metric_section(doc, font, V6_METRICS_KPI)
    page_break(doc)

    # ══ P6 AP ═════════════════════════════════════════════════════════════════
    _annual_eyebrow_title(doc, "05 · ap server", "AP 서버 점검 결과", font)
    add_para(doc, "  · IBM X3650 M4 / Win Server 2012 R2 / UPIS AP 서버.",
             font=font, size=9, color=ANNUAL_GREY_RGB)
    add_para(doc, "")
    _v5_tier_header_table(doc, font, V5_AP_HEADER)
    add_para(doc, "")
    _v5_check_table(doc, font, V5_AP_CHECKS, columns=[
        ("종류",      1.4, "C"),
        ("점검 항목",  3.4, "L"),
        ("점검 방법",  5.4, "L"),
        ("점검 기준",  3.6, "L"),
        ("결과",      1.2, "C"),
        ("메모",      2.4, "L"),
    ])
    _v5_extra_note(doc, font, V5_AP_EXTRA)
    page_break(doc)

    # ══ P7 DB (AIX) ═══════════════════════════════════════════════════════════
    _annual_eyebrow_title(doc, "06 · db server", "DB 서버 (AIX) 점검 결과", font)
    add_para(doc, "  · IBM P720 / AIX 6.1 / UPIS DB 서버 — 육안·구성·성능·DATA·네트워크·프로세서·로그.",
             font=font, size=9, color=ANNUAL_GREY_RGB)
    add_para(doc, "")
    _v5_tier_header_table(doc, font, V5_DB_HEADER)
    add_para(doc, "")
    _v5_check_table(doc, font, V5_DB_CHECKS, columns=[
        ("점검 방법",  1.8, "C"),
        ("점검 항목",  4.0, "L"),
        ("명령어",    6.0, "L"),
        ("결과",      1.2, "C"),
        ("메모",      4.4, "L"),
    ])
    _v5_extra_note(doc, font, V5_DB_EXTRA)
    page_break(doc)

    # ══ P8 DBMS (Oracle 운영 깊이 — V030 시드와 정합) ═══════════════════════════
    _annual_eyebrow_title(doc, "07 · dbms", "DBMS (Oracle) 점검 결과", font)
    add_para(doc, "  · Oracle 운영 점검 — archive · alert · redo · SGA/PGA · tablespace · datafile · 백업 · DG.",
             font=font, size=9, color=ANNUAL_GREY_RGB)
    add_para(doc, "")
    _v5_tier_header_table(doc, font, V5_DBMS_HEADER)
    add_para(doc, "")
    _v5_check_table(doc, font, V5_DBMS_CHECKS, columns=[
        ("구분",       1.4, "C"),
        ("점검 항목",   3.6, "L"),
        ("명령 / SQL", 7.2, "L"),
        ("결과",       1.2, "C"),
        ("메모",       4.0, "L"),
    ])
    _v5_extra_note(doc, font, V5_DBMS_EXTRA)
    page_break(doc)

    # ══ P9 GIS 점검표 (v5 P8 의 표만, 카드는 P10 분할) ════════════════════════
    _annual_eyebrow_title(doc, "08a · gis engine", "GIS 엔진 점검 결과", font)
    add_para(doc, "  · GeoNURIS GSS (Spatial Server) / GWS (GeoWeb Server 64) / Desktop Pro 연계.",
             font=font, size=9, color=ANNUAL_GREY_RGB)
    add_para(doc, "")
    _v5_check_table(doc, font, V5_GIS_CHECKS, columns=[
        ("대상",      2.0, "C"),
        ("점검 항목",  3.6, "L"),
        ("점검 방법",  7.4, "L"),
        ("결과",      1.2, "C"),
        ("메모",      3.2, "L"),
    ])
    _v5_extra_note(doc, font, V5_GIS_EXTRA)
    page_break(doc)

    # ══ P10 GIS 카드 (v6 신규 — GSS/GWS/STORE 3 카드, v4 양식 복원) ═══════════
    _annual_eyebrow_title(doc, "08b · gis cards", "GIS 엔진 상세 분석 (GSS / GWS / Store)", font)
    add_para(doc, "  · 30 일 누적 메트릭 + UWES Store 보존 확인 — agent 자동수집.",
             font=font, size=9, color=ANNUAL_GREY_RGB)
    add_para(doc, "")
    _annual_log_analysis_cards_filled(doc, font, V6_GIS_CARDS)
    page_break(doc)

    # ══ P11 APP UPIS 14 메뉴 ══════════════════════════════════════════════════
    _annual_eyebrow_title(doc, "09 · application", "표준시스템 (UPIS 애플리케이션) 점검 결과", font)
    add_para(doc, "  · 도시계획 / 통계 / 전자심의 / 지구단위계획 / 비정형 / 관리자 / GIS 연동 — 14 개 메뉴.",
             font=font, size=9, color=ANNUAL_GREY_RGB)
    add_para(doc, "")
    _v5_check_table(doc, font, V5_APP_CHECKS, columns=[
        ("분류",        3.0, "C"),
        ("세부 분류",    3.6, "L"),
        ("점검 내용",    7.6, "L"),
        ("결과",        1.2, "C"),
        ("메모",        2.0, "L"),
    ])
    page_break(doc)

    # ══ P12 next round (v5 P10 — 서명 페이지 제거로 마지막) ═══════════════════
    _annual_eyebrow_title(doc, "10 · next round", "차회 점검 계획", font)
    add_para(doc, "")
    for _eyebrow, title, body in FILLED_NEXT_PLAN:
        p = doc.add_paragraph()
        _tight(p, after_pt=4)
        set_run(p.add_run(title),
                font=font, size=13, bold=True, color=ANNUAL_DARK_RGB)
        body_lines = [body] if isinstance(body, str) else list(body)
        for line in body_lines:
            p = doc.add_paragraph()
            _tight(p, after_pt=2)
            set_run(p.add_run(line),
                    font=font, size=10, color=ANNUAL_DARK_RGB)
        add_para(doc, "")
    # ~~v5 P11 signature 제거~~ (갤럭시탭 디지털 서명으로 대체, P1 표지 sig 표는 유지)

    path = OUT_DIR / "점검내역서_시안D_v6.docx"
    doc.save(path)
    return path


# ─────────────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    import sys
    # PowerShell cp949 회피 — stdout 강제 UTF-8
    try:
        sys.stdout.reconfigure(encoding="utf-8")
    except Exception:
        pass

    results = [
        ("D (annual report)", make_annual()),
        ("D-filled (단양 UPIS 5월)", make_annual_filled()),
        ("D-v5 (원본 누락 통합)", make_annual_filled_v5()),
        ("D-v6 (메트릭+카드 분할, 서명 제거)", make_annual_filled_v6()),
    ]
    print("-" * 60)
    for label, path in results:
        size = path.stat().st_size
        print(f"  [OK] draft {label:20s} -> {path.name}  ({size:,} B)")
    print("-" * 60)
    print(f"output dir: {OUT_DIR}")
