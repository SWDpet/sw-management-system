import docx
import sys
import os
sys.stdout.reconfigure(encoding='utf-8')

files = [
    r"C:\Users\PUJ\Downloads\점검내역서_UPIS_SW샘플.docx",
    r"C:\Users\PUJ\Downloads\점검내역서_UPIS_SW_HW샘플.docx",
    r"C:\Users\PUJ\Downloads\점검내역서_KRAS샘플.docx",
]

for fpath in files:
    fname = os.path.basename(fpath)
    print(f"\n{'='*80}")
    print(f"=== {fname} ===")
    print(f"{'='*80}")

    doc = docx.Document(fpath)

    # Extract all tables
    for ti, table in enumerate(doc.tables):
        print(f"\n--- Table {ti+1} ---")
        for ri, row in enumerate(table.rows):
            cells = [cell.text.strip().replace('\n', ' | ') for cell in row.cells]
            # Remove consecutive duplicate cells (merged cells show same text)
            deduped = []
            prev = None
            for c in cells:
                if c != prev:
                    deduped.append(c)
                prev = c
            print(f"  Row {ri}: {' || '.join(deduped)}")

    # Also extract paragraphs (for headers between tables)
    print(f"\n--- Paragraphs ---")
    for pi, para in enumerate(doc.paragraphs):
        if para.text.strip():
            print(f"  P{pi}: [{para.style.name}] {para.text.strip()}")
